#!/usr/bin/env python3

"""PORTFOLIO TERMINAL v4.0 - Bloomberg Style with Backtesting"""

import os, sys

# Add parent directory to path for module imports when running directly
_current_dir = os.path.dirname(os.path.abspath(__file__))
_parent_dir = os.path.dirname(_current_dir)
if _parent_dir not in sys.path:
    sys.path.insert(0, _parent_dir)

import json, webbrowser, threading, math, glob, html, tempfile

from datetime import datetime, timedelta

from http.server import HTTPServer, SimpleHTTPRequestHandler

import urllib.parse

from typing import List, Dict, Any, Optional, Tuple

try:
    import requests
    REQUESTS_OK = True
except ImportError:
    REQUESTS_OK = False

from olyos.logger import get_logger, configure as configure_logging
from olyos.services.api_client import ParallelAPIClient, BatchProgress
from olyos.services.alerts import AlertsService, AlertConfig, create_alerts_service
from olyos.services.benchmark import BenchmarkService, BENCHMARKS, create_benchmark_service
from olyos.services.dividends import DividendsService, create_dividends_service
from olyos.services.position_manager import PositionManager, create_position_manager
from olyos.services.pdf_report import PDFReportService, create_pdf_report_service
from olyos.services.insider import InsiderService, InsiderTransaction, TransactionType, create_insider_service
from olyos.services.rebalancing import RebalancingService, RebalanceConfig, create_rebalancing_service
from olyos.services.ai_analysis import run_analysis as run_ai_analysis
try:
    from olyos.olyos_portfolio_advisor import run_analysis as run_portfolio_advisor_analysis
    PORTFOLIO_ADVISOR_OK = True
except Exception:
    run_portfolio_advisor_analysis = None
    PORTFOLIO_ADVISOR_OK = False

# Initialize loggers for different components
log = get_logger('main')
log_api = get_logger('api')
log_cache = get_logger('cache')
log_backtest = get_logger('backtest')
log_screener = get_logger('screener')
log_portfolio = get_logger('portfolio')



# Data directory (relative to project root)
_DATA_DIR = os.path.join(_parent_dir, 'data')

CONFIG = {
    'portfolio_file': os.path.join(_DATA_DIR, 'portfolio.xlsx'),
    'transactions_file': os.path.join(_DATA_DIR, 'transactions.json'),
    'watchlist_file': os.path.join(_DATA_DIR, 'watchlist.json'),
    'screener_cache_file': os.path.join(_DATA_DIR, 'screener_cache.json'),
    'nav_history_file': os.path.join(_DATA_DIR, 'nav_history.json'),
    'backtest_cache_dir': os.path.join(_DATA_DIR, 'backtest_cache'),
    'backtest_history_file': os.path.join(_DATA_DIR, 'backtest_history.json'),
    'memo_dir': _DATA_DIR,
    'reports_dir': os.path.join(_parent_dir, 'docs', 'reports'),
    'port': 8080,
    'cache_days': 30
}



# EOD Historical Data API

EOD_API_KEY = os.environ.get('EOD_API_KEY')

EOD_OK = EOD_API_KEY is not None and REQUESTS_OK



# Global variable to track data refresh progress

REFRESH_STATUS = {

    'running': False,

    'progress': 0,

    'total': 0,

    'current_ticker': '',

    'message': ''

}

# Global alerts service (initialized lazily)
_ALERTS_SERVICE = None

def get_alerts_service():
    """Get or create the alerts service singleton"""
    global _ALERTS_SERVICE
    if _ALERTS_SERVICE is None:
        _ALERTS_SERVICE = AlertsService(
            watchlist_file=CONFIG['watchlist_file'],
            get_fundamentals_func=eod_get_fundamentals,
            get_prices_func=eod_get_historical_prices
        )
    return _ALERTS_SERVICE

# Global benchmark service (initialized lazily)
_BENCHMARK_SERVICE = None

def get_benchmark_service():
    """Get or create the benchmark service singleton"""
    global _BENCHMARK_SERVICE
    if _BENCHMARK_SERVICE is None:
        benchmark_cache_dir = os.path.join(_DATA_DIR, 'benchmark_cache')
        _BENCHMARK_SERVICE = BenchmarkService(
            cache_dir=benchmark_cache_dir,
            nav_history_file=CONFIG['nav_history_file'],
            get_prices_func=eod_get_historical_prices
        )
    return _BENCHMARK_SERVICE

# Global dividends service (initialized lazily)
_DIVIDENDS_SERVICE = None

def get_dividends_service():
    """Get or create the dividends service singleton"""
    global _DIVIDENDS_SERVICE
    if _DIVIDENDS_SERVICE is None:
        dividends_cache_file = os.path.join(_DATA_DIR, 'dividends_cache.json')
        _DIVIDENDS_SERVICE = DividendsService(
            cache_file=dividends_cache_file,
            get_dividends_func=eod_get_dividends,
            get_fundamentals_func=eod_get_fundamentals
        )
    return _DIVIDENDS_SERVICE

# Global position manager (initialized lazily)
_POSITION_MANAGER = None

def get_position_manager():
    """Get or create the position manager singleton"""
    global _POSITION_MANAGER
    if _POSITION_MANAGER is None:
        def get_current_price(ticker: str) -> float:
            """Get current price for a ticker from cache or API"""
            try:
                fund, err = eod_get_fundamentals(ticker, use_cache=True)
                if fund and not err:
                    price = fund.get('General', {}).get('LastClose')
                    if price:
                        return float(price)
            except:
                pass
            return 0.0

        _POSITION_MANAGER = PositionManager(
            transactions_file=CONFIG['transactions_file'],
            get_price_func=get_current_price
        )
    return _POSITION_MANAGER

# Global PDF report service (initialized lazily)
_PDF_REPORT_SERVICE = None

def get_pdf_report_service():
    """Get or create the PDF report service singleton"""
    global _PDF_REPORT_SERVICE
    if _PDF_REPORT_SERVICE is None:
        # Ensure reports directory exists
        os.makedirs(CONFIG['reports_dir'], exist_ok=True)

        _PDF_REPORT_SERVICE = PDFReportService(
            reports_dir=CONFIG['reports_dir'],
            nav_history_file=CONFIG['nav_history_file'],
            portfolio_func=load_portfolio,
            benchmark_service=get_benchmark_service(),
            position_manager=get_position_manager()
        )
    return _PDF_REPORT_SERVICE

# Global Insider service (initialized lazily)
_INSIDER_SERVICE = None

def get_insider_service():
    """Get or create the Insider service singleton"""
    global _INSIDER_SERVICE
    if _INSIDER_SERVICE is None:
        cache_file = os.path.join(_DATA_DIR, 'insider_cache.json')
        _INSIDER_SERVICE = InsiderService(
            eod_api_key=EOD_API_KEY or '',
            cache_file=cache_file,
            cache_days=3
        )
    return _INSIDER_SERVICE

# Global Rebalancing service (initialized lazily)
_REBALANCING_SERVICE = None

def get_rebalancing_service():
    """Get or create the Rebalancing service singleton"""
    global _REBALANCING_SERVICE
    if _REBALANCING_SERVICE is None:
        _REBALANCING_SERVICE = RebalancingService(
            config=RebalanceConfig(
                max_position_weight=10.0,
                min_position_weight=2.0,
                max_sector_weight=30.0,
                min_higgons_score=4,
                max_pe=17.0
            )
        )
    return _REBALANCING_SERVICE


# Universe of tickers for backtesting

UNIVERSE_FRANCE = []

UNIVERSE_EUROPE = []



def load_universe_from_eod(exchange: str = 'PA', max_market_cap: float = 10e9, min_market_cap: float = 50e6) -> List[str]:

    """Load list of tickers from EOD for a given exchange"""

    if not EOD_OK:

        return []

    

    cache_file = os.path.join(CONFIG['backtest_cache_dir'], f'universe_{exchange}.json')

    os.makedirs(CONFIG['backtest_cache_dir'], exist_ok=True)

    

    # Check cache (valid for 30 days)

    if os.path.exists(cache_file):

        try:

            cache_data = json.load(open(cache_file, encoding='utf-8'))

            cache_date = datetime.strptime(cache_data.get('_cache_date', '2000-01-01'), '%Y-%m-%d')

            if (datetime.now() - cache_date).days < 30:

                return cache_data.get('tickers', [])

        except Exception:
            pass

    

    try:

        # Get list of all tickers on exchange

        url = f"https://eodhd.com/api/exchange-symbol-list/{exchange}?api_token={EOD_API_KEY}&fmt=json"

        log_api.info(f"Loading universe for {exchange}...")

        response = requests.get(url, timeout=60)

        

        if response.status_code != 200:

            log_api.error(f"Error loading universe: {response.status_code}")

            return []

        

        all_tickers = response.json()

        

        # Filter for stocks only (not ETFs, bonds, etc.)

        stocks = [t for t in all_tickers if t.get('Type') == 'Common Stock']

        

        # IMPORTANT: Filter out ADRs, secondary listings, and garbage tickers

        filtered_stocks = []

        for s in stocks:

            code = s.get('Code', '')

            name = s.get('Name', '')

            

            # Skip tickers starting with numbers (usually ADRs on LSE like 0ABC)

            if code and code[0].isdigit():

                continue

            

            # Skip very short tickers (often indices or special instruments)

            if len(code) < 2:

                continue

            

            # Skip tickers with special characters

            if any(c in code for c in ['$', '#', '&', ' ', '=']):

                continue

            

            # Skip common ADR/GDR patterns

            name_upper = name.upper() if name else ''

            if any(x in name_upper for x in ['ADR', 'GDR', 'DEPOSITARY', 'RECEIPT', 'SPONSORED']):

                continue

            

            # Skip ETFs, funds, warrants, etc.

            if any(x in name_upper for x in ['ETF', 'FUND', 'WARRANT', 'TRACKER', 'CERTIFICATE', 'REIT', 'TRUST']):

                continue

            

            filtered_stocks.append(s)

        

        # Build list with basic info

        tickers = []

        for s in filtered_stocks:

            tickers.append({

                'ticker': f"{s['Code']}.{exchange}",

                'name': s.get('Name', ''),

                'exchange': exchange,

                'isin': s.get('Isin', '')

            })

        

        # Save to cache

        cache_data = {

            '_cache_date': datetime.now().strftime('%Y-%m-%d'),

            'tickers': tickers

        }

        with open(cache_file, 'w', encoding='utf-8') as f:

            json.dump(cache_data, f)

        

        log_api.info(f"Found {len(tickers)} stocks on {exchange} (filtered from {len(all_tickers)} total)")

        return tickers

        

    except Exception as e:

        log_api.error(f"Error: {e}")

        return []



def init_universes():

    """Initialize France and Europe universes"""

    global UNIVERSE_FRANCE, UNIVERSE_EUROPE

    

    # France (Euronext Paris)

    UNIVERSE_FRANCE = load_universe_from_eod('PA')

    

    # Europe - major exchanges

    europe_exchanges = ['PA', 'AS', 'BR', 'MI', 'MC', 'XETRA', 'SW', 'LSE']

    

    all_europe = []

    for ex in europe_exchanges:

        tickers = load_universe_from_eod(ex)

        all_europe.extend(tickers)

    

    UNIVERSE_EUROPE = all_europe

    

    log.info(f"Universe loaded: France: {len(UNIVERSE_FRANCE)} tickers, Europe: {len(UNIVERSE_EUROPE)} tickers")



MEMO_PATTERNS = {

    'CARM.PA': ['Memo_Carmila*', 'Investment_Memo_Carmila*'],

    'MAU.PA': ['Memo_Maurel*', 'Investment_Memo_Maurel*'],

    'STF.PA': ['Memo_STEF*', 'Investment_Memo_STEF*', 'Memo_Stef*'],

    'ALCAT.PA': ['Memo_Catana*', 'Investment_Memo_Catana*'],

    'GTT.PA': ['Memo_GTT*', 'Investment_Memo_GTT*'],

    'ARG.PA': ['Memo_Argan*', 'Investment_Memo_ARGAN*'],

    'ICAD.PA': ['Memo_Icade*', 'Investment_Memo_ICADE*'],

    'ALREW.PA': ['Memo_Reworld*', 'Investment_Memo_Reworld*'],

    'ALHOP.PA': ['Memo_Hopscotch*', 'Investment_Memo_HOPSCOTCH*'],

    'STF': ['Memo_STEF*', 'Memo_Stef*'],

    'ALREW': ['Memo_Reworld*'],

    'ALHOP': ['Memo_Hopscotch*', 'Investment_Memo_HOPSCOTCH*'],

}



EUROPE_DB = [

    {"ticker": "VCT.PA", "name": "Vicat", "sector": "Materiaux", "country": "France"},

    {"ticker": "FGR.PA", "name": "Eiffage", "sector": "Construction", "country": "France"},

    {"ticker": "SPIE.PA", "name": "Spie", "sector": "Services", "country": "France"},

    {"ticker": "GTT.PA", "name": "GTT", "sector": "Energie", "country": "France"},

    {"ticker": "TRI.PA", "name": "Trigano", "sector": "Automobile", "country": "France"},

    {"ticker": "BEN.PA", "name": "Beneteau", "sector": "Nautisme", "country": "France"},

    {"ticker": "ELIS.PA", "name": "Elis", "sector": "Services", "country": "France"},

    {"ticker": "STF.PA", "name": "Stef", "sector": "Transport", "country": "France"},

    {"ticker": "DBG.PA", "name": "Derichebourg", "sector": "Services", "country": "France"},

    {"ticker": "AUB.PA", "name": "Aubay", "sector": "IT", "country": "France"},

    {"ticker": "ARG.PA", "name": "Argan", "sector": "Immobilier", "country": "France"},

    {"ticker": "CARM.PA", "name": "Carmila", "sector": "Immobilier", "country": "France"},

    {"ticker": "ICAD.PA", "name": "Icade", "sector": "Immobilier", "country": "France"},

    {"ticker": "ALREW.PA", "name": "Reworld Media", "sector": "Medias", "country": "France"},

    {"ticker": "ALHOP.PA", "name": "Hopscotch", "sector": "Communication", "country": "France"},

    {"ticker": "FNAC.PA", "name": "Fnac Darty", "sector": "Distribution", "country": "France"},

    {"ticker": "MAU.PA", "name": "Maurel Prom", "sector": "Energie", "country": "France"},

    {"ticker": "RUI.PA", "name": "Rubis", "sector": "Energie", "country": "France"},

    {"ticker": "VK.PA", "name": "Vallourec", "sector": "Energie", "country": "France"},

    {"ticker": "TE.PA", "name": "Technip Energies", "sector": "Energie", "country": "France"},

    {"ticker": "NEX.PA", "name": "Nexans", "sector": "Industrie", "country": "France"},

    {"ticker": "SK.PA", "name": "SEB", "sector": "Consommation", "country": "France"},

    {"ticker": "ERA.PA", "name": "Eramet", "sector": "Materiaux", "country": "France"},

    {"ticker": "COFA.PA", "name": "Coface", "sector": "Finance", "country": "France"},

    {"ticker": "SW.PA", "name": "Sodexo", "sector": "Services", "country": "France"},

    {"ticker": "RI.PA", "name": "Pernod Ricard", "sector": "Consommation", "country": "France"},

    {"ticker": "TTE.PA", "name": "TotalEnergies", "sector": "Energie", "country": "France"},

    {"ticker": "WBD.MI", "name": "Webuild", "sector": "Construction", "country": "Italie"},

    {"ticker": "DAN.MI", "name": "Danieli", "sector": "Industrie", "country": "Italie"},

    {"ticker": "MAIRE.MI", "name": "Maire Tecnimont", "sector": "Ingenierie", "country": "Italie"},

    {"ticker": "BZU.MI", "name": "Buzzi", "sector": "Materiaux", "country": "Italie"},

    {"ticker": "CEM.MI", "name": "Cementir", "sector": "Materiaux", "country": "Italie"},

    {"ticker": "CAF.MC", "name": "CAF", "sector": "Transport", "country": "Espagne"},

    {"ticker": "TRE.MC", "name": "Tecnicas Reunidas", "sector": "Ingenierie", "country": "Espagne"},

    {"ticker": "IDR.MC", "name": "Indra Sistemas", "sector": "IT", "country": "Espagne"},

    {"ticker": "CIE.MC", "name": "CIE Automotive", "sector": "Automobile", "country": "Espagne"},

    {"ticker": "SCYR.MC", "name": "Sacyr", "sector": "Construction", "country": "Espagne"},

    {"ticker": "BAMNB.AS", "name": "Royal BAM", "sector": "Construction", "country": "Pays-Bas"},

    {"ticker": "HEIJM.AS", "name": "Heijmans", "sector": "Construction", "country": "Pays-Bas"},

    {"ticker": "HBH.DE", "name": "Hornbach", "sector": "Distribution", "country": "Allemagne"},

    {"ticker": "WAC.DE", "name": "Wacker Neuson", "sector": "Industrie", "country": "Allemagne"},

    {"ticker": "NDA.DE", "name": "Aurubis", "sector": "Materiaux", "country": "Allemagne"},

    {"ticker": "SZG.DE", "name": "Salzgitter", "sector": "Materiaux", "country": "Allemagne"},

    {"ticker": "BEKB.BR", "name": "Bekaert", "sector": "Industrie", "country": "Belgique"},

    {"ticker": "SIP.BR", "name": "Sipef", "sector": "Agriculture", "country": "Belgique"},

    {"ticker": "KLR.L", "name": "Keller Group", "sector": "Construction", "country": "UK"},

    {"ticker": "MGNS.L", "name": "Morgan Sindall", "sector": "Construction", "country": "UK"},

    {"ticker": "IMB.L", "name": "Imperial Brands", "sector": "Consommation", "country": "UK"},

    {"ticker": "METLEN.AT", "name": "Metlen Energy", "sector": "Energie", "country": "Grece"},

    {"ticker": "MOH.AT", "name": "Motor Oil", "sector": "Energie", "country": "Grece"},

    {"ticker": "BELA.AT", "name": "Jumbo", "sector": "Distribution", "country": "Grece"},

    {"ticker": "IMPN.SW", "name": "Implenia", "sector": "Construction", "country": "Suisse"},

    {"ticker": "EGL.LS", "name": "Mota-Engil", "sector": "Construction", "country": "Portugal"},

    {"ticker": "POS.VI", "name": "PORR", "sector": "Construction", "country": "Autriche"},

    {"ticker": "PNDORA.CO", "name": "Pandora", "sector": "Luxe", "country": "Danemark"},

    {"ticker": "ALCAT.PA", "name": "Catana Group", "sector": "Nautisme", "country": "France"},

]



try:

    import pandas as pd

    PANDAS_OK = True

except Exception:
    PANDAS_OK = False



try:

    import yfinance as yf

    YFINANCE_OK = True

except Exception:
    YFINANCE_OK = False



try:

    import anthropic

    ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')

    ANTHROPIC_OK = ANTHROPIC_API_KEY is not None

except Exception:
    ANTHROPIC_OK = False

    ANTHROPIC_API_KEY = None



def load_json(f, default):

    try:

        return json.load(open(f, encoding='utf-8')) if os.path.exists(f) else default

    except Exception:
        return default



def save_json(f, data):

    json.dump(data, open(f, 'w', encoding='utf-8'), indent=2)



def read_memo_content(filepath):

    """Read content from a .docx memo file and return as styled HTML"""

    if not filepath or not os.path.exists(filepath):

        return None

    

    try:

        from docx import Document

        from docx.table import Table

        doc = Document(filepath)

        

        html_parts = []

        

        for element in doc.element.body:

            # Handle tables

            if element.tag.endswith('tbl'):

                for table in doc.tables:

                    if table._element == element:

                        html_parts.append(render_table_html(table))

                        break

            # Handle paragraphs

            elif element.tag.endswith('p'):

                for para in doc.paragraphs:

                    if para._element == element:

                        html = render_paragraph_html(para)

                        if html:

                            html_parts.append(html)

                        break

        

        return '\n'.join(html_parts) if html_parts else None

        

    except ImportError:

        # python-docx not installed, try basic extraction

        try:

            import zipfile

            import xml.etree.ElementTree as ET

            

            with zipfile.ZipFile(filepath) as z:

                xml_content = z.read('word/document.xml')

            

            tree = ET.fromstring(xml_content)

            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            

            paragraphs = []

            for para in tree.findall('.//w:p', ns):

                texts = [node.text for node in para.findall('.//w:t', ns) if node.text]

                if texts:

                    full_text = ''.join(texts).strip()

                    if full_text:

                        # Detect headers by keywords

                        if any(kw in full_text.upper() for kw in ['SIGNAL', 'RESUME', 'SUMMARY', 'THESIS', 'RISQUES', 'VALORISATION', 'FINANCIER', 'CONCLUSION', 'PROFIL', 'ACTIONNARIAT', 'GEOGRAPHIE']):

                            paragraphs.append(f'<div class="memo-section-title">{full_text}</div>')

                        else:

                            paragraphs.append(f'<p class="memo-text">{full_text}</p>')

            

            return '\n'.join(paragraphs) if paragraphs else None

            

        except Exception as e:

            log.error(f"Error reading memo: {e}")

            return None

    except Exception as e:

        log.error(f"Error reading memo: {e}")

        return None





def render_paragraph_html(para):

    """Convert a docx paragraph to styled HTML"""

    text = para.text.strip()

    if not text:

        return None

    

    # Detect style

    style_name = para.style.name if para.style else ''

    

    # Check if bold (title/header)

    is_bold = False

    is_all_caps = text.isupper() and len(text) > 3

    if para.runs:

        is_bold = para.runs[0].bold

    

    # Section titles (detect by keywords or style)

    section_keywords = ['SIGNAL', 'RESUME', 'SUMMARY', 'THESIS', 'RISQUES', 'RISK', 'VALORISATION', 

                        'VALUATION', 'FINANCIER', 'FINANCIAL', 'CONCLUSION', 'PROFIL', 'PROFILE',

                        'ACTIONNARIAT', 'SHAREHOLDERS', 'GEOGRAPHIE', 'GEOGRAPHY', 'CATALYSTS',

                        'CATALYSEURS', 'POINTS FORTS', 'STRENGTHS', 'POINTS FAIBLES', 'WEAKNESSES',

                        'RECOMMENDATION', 'RECOMMANDATION', 'OVERVIEW', 'APERÃ‡U', 'DESCRIPTION',

                        'STRUCTURE', 'STRATEGIE', 'STRATEGY', 'HISTORIQUE', 'HISTORY', 'ACQUISITION']

    

    is_section_title = any(kw in text.upper() for kw in section_keywords) or 'Heading' in style_name

    

    # Signal/Verdict detection

    signal_keywords = {'ACHAT': 'buy', 'BUY': 'buy', 'STRONG BUY': 'buy', 

                       'VENTE': 'sell', 'SELL': 'sell', 'ECARTER': 'sell',

                       'SURVEILLANCE': 'watch', 'WATCH': 'watch', 'HOLD': 'hold',

                       'NEUTRE': 'hold', 'NEUTRAL': 'hold'}

    

    for kw, cls in signal_keywords.items():

        if kw in text.upper():

            return f'<div class="memo-signal memo-signal-{cls}">{text}</div>'

    

    # Format based on type

    if is_section_title or is_all_caps:

        return f'<div class="memo-section-title">{text}</div>'

    elif is_bold:

        return f'<div class="memo-subtitle">{text}</div>'

    elif text.startswith(' - ') or text.startswith('-') or text.startswith(' - '):

        return f'<div class="memo-bullet">{text}</div>'

    elif ':' in text and len(text.split(':')[0]) < 30:

        # Key-value pair

        parts = text.split(':', 1)

        return f'<div class="memo-kv"><span class="memo-key">{parts[0]}:</span><span class="memo-value">{parts[1].strip()}</span></div>'

    else:

        return f'<p class="memo-text">{text}</p>'





def render_table_html(table):

    """Convert a docx table to styled HTML"""

    html = ['<table class="memo-table">']

    

    for i, row in enumerate(table.rows):

        if i == 0:

            html.append('<thead><tr>')

            for cell in row.cells:

                html.append(f'<th>{cell.text.strip()}</th>')

            html.append('</tr></thead><tbody>')

        else:

            html.append('<tr>')

            for j, cell in enumerate(row.cells):

                cell_text = cell.text.strip()

                # Detect numeric values for alignment

                cell_class = 'num' if any(c.isdigit() for c in cell_text) and j > 0 else ''

                # Detect positive/negative

                if cell_class == 'num':

                    if cell_text.startswith('+') or (cell_text.replace(',', '').replace('.', '').replace('%', '').replace(' ', '').isdigit() and not cell_text.startswith('-')):

                        cell_class += ' pos'

                    elif cell_text.startswith('-'):

                        cell_class += ' neg'

                html.append(f'<td class="{cell_class}">{cell_text}</td>')

            html.append('</tr>')

    

    html.append('</tbody></table>')

    return '\n'.join(html)



def get_yf_ticker(ticker):

    """Convert ticker to Yahoo Finance format"""

    # Already has a suffix

    if '.' in ticker:

        return ticker

    

    # Special tickers mapping (non-French or special cases)

    TICKER_MAP = {

        # Belgian stocks

        'SIP': 'SIP.BR',           # Sipef

        'BEKB': 'BEKB.BR',         # Bekaert

        

        # Danish stocks  

        'PNDORA': 'PNDORA.CO',     # Pandora

        

        # German stocks

        'NEAG': 'NEAG.DE',         # Nagarro

        'HBH': 'HBH.DE',           # Hornbach

        'WAC': 'WAC.DE',           # Wacker Neuson

        'NDA': 'NDA.DE',           # Aurubis

        'SZG': 'SZG.DE',           # Salzgitter

        

        # Swiss stocks

        'ZURN': 'ZURN.SW',         # Zurich Insurance

        'IMPN': 'IMPN.SW',         # Implenia

        

        # US stocks (no suffix needed)

        'FCX': 'FCX',              # Freeport McMoran

        

        # Dutch stocks

        'BAMNB': 'BAMNB.AS',       # Royal BAM

        'HEIJM': 'HEIJM.AS',       # Heijmans

        

        # Italian stocks

        'WBD': 'WBD.MI',           # Webuild

        'DAN': 'DAN.MI',           # Danieli

        'MAIRE': 'MAIRE.MI',       # Maire Tecnimont

        'BZU': 'BZU.MI',           # Buzzi

        'CEM': 'CEM.MI',           # Cementir

        

        # Spanish stocks

        'CAF': 'CAF.MC',           # CAF

        'TRE': 'TRE.MC',           # Tecnicas Reunidas

        'IDR': 'IDR.MC',           # Indra

        'CIE': 'CIE.MC',           # CIE Automotive

        'SCYR': 'SCYR.MC',         # Sacyr

        

        # Greek stocks

        'METLEN': 'METLEN.AT',     # Metlen

        'MOH': 'MOH.AT',           # Motor Oil

        'BELA': 'BELA.AT',         # Jumbo

        

        # Portuguese stocks

        'EGL': 'EGL.LS',           # Mota-Engil

        

        # Austrian stocks

        'POS': 'POS.VI',           # PORR

        

        # UK stocks

        'KLR': 'KLR.L',            # Keller

        'MGNS': 'MGNS.L',          # Morgan Sindall

        'IMB': 'IMB.L',            # Imperial Brands

        'RIO': 'RIO.L',            # Rio Tinto

        'BA': 'BA.L',              # BAE Systems

        

        # French stocks (explicit mapping for safety)

        'MC': 'MC.PA',             # LVMH

        'VIE': 'VIE.PA',           # Veolia

        'GFC': 'GFC.PA',           # Gecina

        'ALVAP': 'ALVAP.PA',       # Valneva

        'ALWEC': 'ALWEC.PA',       # We.Connect

        'NXI': 'NXI.PA',           # Nexity

        'SK': 'SK.PA',             # SEB

        'ALCAT': 'ALCAT.PA',       # Catana Group

        'CATG': 'ALCAT.PA',        # Catana Group (old ticker)

    }

    

    ticker_upper = ticker.upper()

    if ticker_upper in TICKER_MAP:

        return TICKER_MAP[ticker_upper]

    

    # Default: assume French stock

    return ticker + '.PA'



def find_memo(ticker):

    """Find memo file for a given ticker"""

    # Try with and without .PA suffix

    tickers_to_try = [ticker, ticker.replace('.PA', ''), ticker.split('.')[0]]

    

    for t in tickers_to_try:

        patterns = MEMO_PATTERNS.get(t, [])

        for pattern in patterns:

            matches = glob.glob(os.path.join(CONFIG['memo_dir'], pattern))

            if matches:

                return matches[0]

    

    # Generic patterns

    ticker_base = ticker.split('.')[0]

    generic_patterns = [f'Memo_{ticker_base}*', f'Investment_Memo_{ticker_base}*', f'Memo_*{ticker_base}*']

    for pattern in generic_patterns:

        matches = glob.glob(os.path.join(CONFIG['memo_dir'], pattern))

        if matches:

            return matches[0]

    return None



def load_watchlist() -> List[str]:

    return load_json(CONFIG['watchlist_file'], [])



def save_watchlist(w: List[str]) -> None:

    save_json(CONFIG['watchlist_file'], w)



# =============================================================================

# EOD HISTORICAL DATA API FUNCTIONS + SMART CACHING

# =============================================================================



# Cache settings

CACHE_DIR = 'backtest_cache'

FUNDAMENTALS_CACHE_DAYS = 30  # Refresh fundamentals monthly

PRICES_CACHE_DAYS = 1  # Refresh prices daily (only if online)

UNIVERSE_CACHE_DAYS = 30  # Refresh universe monthly



def ensure_cache_dir():

    """Ensure cache directory exists"""

    os.makedirs(CACHE_DIR, exist_ok=True)

    os.makedirs(os.path.join(CACHE_DIR, 'fundamentals'), exist_ok=True)

    os.makedirs(os.path.join(CACHE_DIR, 'prices'), exist_ok=True)

    os.makedirs(os.path.join(CACHE_DIR, 'universe'), exist_ok=True)

    os.makedirs(os.path.join(CACHE_DIR, 'dividends'), exist_ok=True)



def get_cache_path(category, key):

    """Get cache file path for a given category and key"""

    ensure_cache_dir()

    safe_key = key.replace('.', '_').replace('^', '_').replace('/', '_')

    return os.path.join(CACHE_DIR, category, f'{safe_key}.json')



def is_cache_valid(cache_file, max_days):

    """Check if cache file exists and is not too old"""

    if not os.path.exists(cache_file):

        return False

    try:

        with open(cache_file, 'r', encoding='utf-8') as f:

            data = json.load(f)

        cache_date = datetime.strptime(data.get('_cache_date', '2000-01-01'), '%Y-%m-%d')

        return (datetime.now() - cache_date).days < max_days

    except Exception:
        return False



def load_from_cache(cache_file):

    """Load data from cache file"""

    try:

        with open(cache_file, 'r', encoding='utf-8') as f:

            return json.load(f)

    except Exception:
        return None



def save_to_cache(cache_file, data):

    """Save data to cache file with timestamp"""

    ensure_cache_dir()

    data['_cache_date'] = datetime.now().strftime('%Y-%m-%d')

    with open(cache_file, 'w', encoding='utf-8') as f:

        json.dump(data, f)



def get_cache_stats():

    """Get statistics about cached data"""

    ensure_cache_dir()

    stats = {

        'fundamentals': 0,

        'prices': 0,

        'universe': 0,

        'total_size_mb': 0

    }

    

    for category in ['fundamentals', 'prices', 'universe']:

        cat_path = os.path.join(CACHE_DIR, category)

        if os.path.exists(cat_path):

            files = [f for f in os.listdir(cat_path) if f.endswith('.json')]

            stats[category] = len(files)

            for f in files:

                stats['total_size_mb'] += os.path.getsize(os.path.join(cat_path, f)) / (1024*1024)

    

    stats['total_size_mb'] = round(stats['total_size_mb'], 2)

    return stats



# =============================================================================

# BACKTEST HISTORY MANAGEMENT

# =============================================================================



def load_backtest_history():

    """Load saved backtests history"""

    return load_json(CONFIG['backtest_history_file'], [])



def save_backtest_history(history):

    """Save backtests history"""

    log_backtest.info(f"Saving {len(history)} backtests to {CONFIG['backtest_history_file']}")

    save_json(CONFIG['backtest_history_file'], history)



def save_backtest_result(results, name=None):

    """Save a backtest result to history"""

    history = load_backtest_history()

    

    # Generate name if not provided

    if not name:

        params = results.get('params', {})

        scope = params.get('universe_scope', 'custom')

        pe = params.get('pe_max', '?')

        roe = params.get('roe_min', '?')

        name = f"{scope.upper()} PE<={pe} ROE>={roe}%"

    

    # Create summary (don't save full trade history to keep file small)

    summary = {

        'id': datetime.now().strftime('%Y%m%d_%H%M%S'),

        'name': name,

        'date': datetime.now().strftime('%Y-%m-%d %H:%M'),

        'params': results.get('params', {}),

        'metrics': results.get('metrics', {}),

        'yearly_returns': results.get('yearly_returns', []),

        'equity_curve_summary': {

            'start': results.get('equity_curve', [{}])[0] if results.get('equity_curve') else {},

            'end': results.get('equity_curve', [{}])[-1] if results.get('equity_curve') else {},

            'points': len(results.get('equity_curve', []))

        },

        'trades_count': len(results.get('trades', [])),

        'errors_count': len(results.get('errors', []))

    }

    

    # Add to history (keep last 50 backtests)

    history.insert(0, summary)

    history = history[:50]

    

    save_backtest_history(history)

    return summary['id']



def delete_backtest(backtest_id):

    """Delete a backtest from history"""

    history = load_backtest_history()

    history = [h for h in history if h.get('id') != backtest_id]

    save_backtest_history(history)



def rename_backtest(backtest_id, new_name):

    """Rename a backtest"""

    history = load_backtest_history()

    for h in history:

        if h.get('id') == backtest_id:

            h['name'] = new_name

            break

    save_backtest_history(history)



# =============================================================================

# AI OPTIMIZER - Find optimal investment criteria

# =============================================================================



def run_ai_optimization(scope='france', optimization_goal='balanced'):

    """

    Use Claude AI to find optimal investment criteria through iterative backtesting

    

    optimization_goal: 'max_return', 'max_sharpe', 'min_drawdown', 'balanced'

    """

    

    if not ANTHROPIC_OK:

        return {'error': 'Anthropic API not configured. Set ANTHROPIC_API_KEY.'}

    

    log_backtest.info("AI OPTIMIZER: Starting optimization...")

    log_backtest.info(f"Scope: {scope}")

    log_backtest.info(f"Goal: {optimization_goal}")

    

    results = {

        'iterations': [],

        'best_params': None,

        'best_metrics': None,

        'ai_analysis': '',

        'recommendation': ''

    }

    

    # Phase 1: Grid search with key parameter combinations

    log_backtest.info("AI OPTIMIZER: Phase 1: Running grid search...")

    

    param_grid = [

        # PE variations

        {'pe_max': 8, 'roe_min': 10, 'pe_sell': 15, 'debt_equity_max': 100, 'max_positions': 20},

        {'pe_max': 10, 'roe_min': 10, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 20},

        {'pe_max': 12, 'roe_min': 10, 'pe_sell': 20, 'debt_equity_max': 100, 'max_positions': 20},

        {'pe_max': 15, 'roe_min': 10, 'pe_sell': 25, 'debt_equity_max': 100, 'max_positions': 20},

        

        # ROE variations

        {'pe_max': 12, 'roe_min': 8, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 20},

        {'pe_max': 12, 'roe_min': 12, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 20},

        {'pe_max': 12, 'roe_min': 15, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 20},

        

        # Debt variations

        {'pe_max': 12, 'roe_min': 10, 'pe_sell': 17, 'debt_equity_max': 50, 'max_positions': 20},

        {'pe_max': 12, 'roe_min': 10, 'pe_sell': 17, 'debt_equity_max': 150, 'max_positions': 20},

        

        # Position count variations

        {'pe_max': 12, 'roe_min': 10, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 10},

        {'pe_max': 12, 'roe_min': 10, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 15},

        {'pe_max': 12, 'roe_min': 10, 'pe_sell': 17, 'debt_equity_max': 100, 'max_positions': 30},

        

        # Combined aggressive value

        {'pe_max': 8, 'roe_min': 12, 'pe_sell': 12, 'debt_equity_max': 50, 'max_positions': 15},

        

        # Combined moderate

        {'pe_max': 15, 'roe_min': 8, 'pe_sell': 25, 'debt_equity_max': 150, 'max_positions': 25},

    ]

    

    grid_results = []

    

    for i, params in enumerate(param_grid):

        log_backtest.info(f"[{i+1}/{len(param_grid)}] Testing PE<={params['pe_max']}, ROE>={params['roe_min']}%, {params['max_positions']} positions...")

        

        backtest_params = {

            'start_date': '2014-01-01',

            'end_date': datetime.now().strftime('%Y-%m-%d'),

            'universe_scope': scope,

            'universe': [],

            'pe_max': params['pe_max'],

            'roe_min': params['roe_min'],

            'pe_sell': params['pe_sell'],

            'roe_min_hold': 8,

            'debt_equity_max': params['debt_equity_max'],

            'rebalance_freq': 'quarterly',

            'initial_capital': 100000,

            'max_positions': params['max_positions'],

            'benchmark': '^FCHI'

        }

        

        try:

            bt_result = run_backtest(backtest_params)

            metrics = bt_result.get('metrics', {})

            

            grid_results.append({

                'params': params,

                'metrics': {

                    'total_return': metrics.get('total_return', 0),

                    'cagr': metrics.get('cagr', 0),

                    'sharpe': metrics.get('sharpe', 0),

                    'max_drawdown': metrics.get('max_drawdown', 0),

                    'win_rate': metrics.get('win_rate', 0),

                    'total_trades': metrics.get('total_trades', 0),

                    'alpha': metrics.get('alpha', 0)

                },

                'errors_count': len(bt_result.get('errors', []))

            })

        except Exception as e:

            log_backtest.error(f"Error: {e}")

            continue

    

    results['iterations'] = grid_results

    

    if not grid_results:

        return {'error': 'No successful backtests completed'}

    

    # Phase 2: Send results to Claude for analysis

    log_backtest.info("AI OPTIMIZER: Phase 2: AI Analysis...")

    

    try:

        import anthropic

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        

        # Prepare data summary for Claude

        results_summary = "BACKTEST RESULTS GRID:\n\n"

        for i, r in enumerate(grid_results):

            p = r['params']

            m = r['metrics']

            results_summary += f"""Test {i+1}: PE<={p['pe_max']}, ROE>={p['roe_min']}%, Debt<={p['debt_equity_max']}%, {p['max_positions']} positions, Sell@PE>{p['pe_sell']}

  â†’ Return: {m['total_return']:.1f}%, CAGR: {m['cagr']:.2f}%, Sharpe: {m['sharpe']:.2f}, MaxDD: -{m['max_drawdown']:.1f}%, WinRate: {m['win_rate']:.0f}%, Trades: {m['total_trades']}, Alpha: {m['alpha']:.1f}%



"""

        

        optimization_prompt = f"""Tu es un expert en investissement value et en analyse quantitative. Tu analyses les resultats de backtests d'une strategie d'investissement style "William Higgons" sur les small/mid caps {scope.upper()}.



{results_summary}



OBJECTIF D'OPTIMISATION: {optimization_goal}

- max_return: Maximiser le rendement total

- max_sharpe: Maximiser le ratio de Sharpe (rendement ajuste du risque)

- min_drawdown: Minimiser le drawdown maximum

- balanced: Trouver le meilleur equilibre rendement/risque



ANALYSE DEMANDEE:



1. ANALYSE DES RESULTATS

Analyse les patterns dans les donnees:

- Quel impact a le PE max sur les rendements?

- Quel impact a le ROE min sur les rendements?

- Quel impact a le nombre de positions?

- Y a-t-il des combinaisons clairement superieures?



2. PARAMÃˆTRES OPTIMAUX

Base sur ton analyse et l'objectif "{optimization_goal}", recommande les parametres optimaux:

- PE_MAX: (valeur entre 5 et 20)

- ROE_MIN: (valeur entre 5 et 20)

- PE_SELL: (valeur entre 12 et 30)

- DEBT_EQUITY_MAX: (valeur entre 30 et 200)

- MAX_POSITIONS: (valeur entre 8 et 30)



3. EXPLICATION

Explique pourquoi ces parametres sont optimaux et quels compromis tu as faits.



4. AVERTISSEMENTS

Mentionne les limites de cette optimisation (survivorship bias, periode testee, etc.)



Reponds en JSON avec cette structure exacte:

{{

  "analysis": "ton analyse detaillee",

  "optimal_params": {{

    "pe_max": X,

    "roe_min": X,

    "pe_sell": X,

    "debt_equity_max": X,

    "max_positions": X

  }},

  "expected_metrics": {{

    "cagr_estimate": "X-Y%",

    "sharpe_estimate": "X-Y",

    "max_drawdown_estimate": "X-Y%"

  }},

  "explanation": "explication des choix",

  "warnings": ["warning1", "warning2"],

  "confidence": "high/medium/low"

}}"""



        response = client.messages.create(

            model="claude-sonnet-4-20250514",

            max_tokens=4096,

            messages=[{"role": "user", "content": optimization_prompt}]

        )

        

        ai_response = response.content[0].text

        

        # Parse JSON response

        import re

        json_match = re.search(r'\{[\s\S]*\}', ai_response)

        if json_match:

            ai_data = json.loads(json_match.group())

            results['ai_analysis'] = ai_data.get('analysis', '')

            results['best_params'] = ai_data.get('optimal_params', {})

            results['expected_metrics'] = ai_data.get('expected_metrics', {})

            results['explanation'] = ai_data.get('explanation', '')

            results['warnings'] = ai_data.get('warnings', [])

            results['confidence'] = ai_data.get('confidence', 'medium')

        else:

            results['ai_analysis'] = ai_response

            results['best_params'] = find_best_from_grid(grid_results, optimization_goal)

        

        # Phase 3: Run final backtest with optimal params

        if results['best_params']:

            log_backtest.info("AI OPTIMIZER: Phase 3: Running optimal backtest...")

            

            optimal_bt_params = {

                'start_date': '2014-01-01',

                'end_date': datetime.now().strftime('%Y-%m-%d'),

                'universe_scope': scope,

                'universe': [],

                'pe_max': results['best_params'].get('pe_max', 12),

                'roe_min': results['best_params'].get('roe_min', 10),

                'pe_sell': results['best_params'].get('pe_sell', 17),

                'roe_min_hold': 8,

                'debt_equity_max': results['best_params'].get('debt_equity_max', 100),

                'rebalance_freq': 'quarterly',

                'initial_capital': 100000,

                'max_positions': results['best_params'].get('max_positions', 20),

                'benchmark': '^FCHI'

            }

            

            optimal_result = run_backtest(optimal_bt_params)

            results['best_metrics'] = optimal_result.get('metrics', {})

            results['best_equity_curve'] = optimal_result.get('equity_curve', [])

            results['best_yearly_returns'] = optimal_result.get('yearly_returns', [])

            

            # Save to history

            save_backtest_result(optimal_result, f"🤖 AI OPTIMAL ({scope})")

        

        log_backtest.info("AI OPTIMIZER: Optimization complete!")

        return results

        

    except Exception as e:

        log_backtest.error(f"AI OPTIMIZER: Error in AI analysis: {e}")

        # Fallback: find best from grid without AI

        results['best_params'] = find_best_from_grid(grid_results, optimization_goal)

        results['ai_analysis'] = f"AI analysis failed: {e}. Using grid search best result."

        return results





def find_best_from_grid(grid_results, goal):

    """Find best parameters from grid search results"""

    if not grid_results:

        return None

    

    if goal == 'max_return':

        best = max(grid_results, key=lambda x: x['metrics'].get('total_return', 0))

    elif goal == 'max_sharpe':

        best = max(grid_results, key=lambda x: x['metrics'].get('sharpe', 0))

    elif goal == 'min_drawdown':

        best = min(grid_results, key=lambda x: x['metrics'].get('max_drawdown', 100))

    else:  # balanced

        # Score = CAGR * Sharpe / (MaxDD + 1)

        def score(r):

            m = r['metrics']

            cagr = m.get('cagr', 0)

            sharpe = max(m.get('sharpe', 0), 0.01)

            dd = m.get('max_drawdown', 50)

            return (cagr * sharpe) / (dd + 1) if dd > 0 else cagr * sharpe

        

        best = max(grid_results, key=score)

    

    return best['params']



def get_eod_ticker(ticker):

    """Convert ticker to EOD format"""

    

    # Pour les benchmarks, utiliser des ETFs qui trackent les indices

    # Car les indices purs (.INDX) ne sont pas toujours disponibles sur EOD

    BENCHMARK_ETF_MAP = {

        # Utiliser des ETFs populaires qui trackent les indices

        '^FCHI': 'CAC.PA',             # Lyxor CAC 40 ETF

        '^STOXX50E': 'MSE.PA',         # Lyxor Euro Stoxx 50 ETF  

        '^STOXX': 'MEUD.PA',           # Lyxor Stoxx Europe 600 ETF

        '^GDAXI': 'DAX.PA',            # Lyxor DAX ETF

        '^FTSE': 'VUKE.LSE',           # Vanguard FTSE 100 ETF

        '^AEX': 'IAEX.AS',             # iShares AEX ETF

    }

    

    if ticker in BENCHMARK_ETF_MAP:

        return BENCHMARK_ETF_MAP[ticker]

    

    if '.' in ticker:

        return ticker

    return ticker + '.PA'



def eod_get_fundamentals(ticker: str, use_cache: bool = True, force_refresh: bool = False) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:

    """Get fundamental data from EOD Historical Data API with smart caching"""

    cache_file = get_cache_path('fundamentals', ticker)

    

    # Try cache first (unless force refresh)

    if use_cache and not force_refresh and is_cache_valid(cache_file, FUNDAMENTALS_CACHE_DAYS):

        cached = load_from_cache(cache_file)

        if cached:

            return cached, None

    

    # If no API key, try to use cache even if expired

    if not EOD_OK:

        cached = load_from_cache(cache_file)

        if cached:

            log_cache.info(f"Using offline cache for {ticker} fundamentals")

            return cached, None

        return None, "EOD API not configured and no cached data available"

    

    try:

        eod_ticker = get_eod_ticker(ticker)

        url = f"https://eodhd.com/api/fundamentals/{eod_ticker}?api_token={EOD_API_KEY}&fmt=json"

        log_api.info(f"Fetching fundamentals: {ticker}")

        response = requests.get(url, timeout=30)

        

        if response.status_code == 200:

            data = response.json()

            save_to_cache(cache_file, data)

            return data, None

        else:

            # Try cache on API error

            cached = load_from_cache(cache_file)

            if cached:

                log_cache.info(f"API error, using cached data for {ticker}")

                return cached, None

            return None, f"API error: {response.status_code}"

            

    except Exception as e:

        # Try cache on exception

        cached = load_from_cache(cache_file)

        if cached:

            log_cache.info(f"Exception, using cached data for {ticker}")

            return cached, None

        return None, str(e)



def eod_get_historical_prices(ticker: str, start_date: str, end_date: str, use_cache: bool = True, force_refresh: bool = False) -> Tuple[Optional[List[Dict[str, Any]]], Optional[str]]:

    """Get historical prices from EOD with smart caching"""

    # Create cache key that includes date range

    cache_key = f"{ticker}_{start_date}_{end_date}"

    cache_file = get_cache_path('prices', cache_key)

    

    # Try cache first

    if use_cache and not force_refresh and is_cache_valid(cache_file, PRICES_CACHE_DAYS):

        cached = load_from_cache(cache_file)

        if cached and cached.get('prices'):

            return cached['prices'], None

    

    # For historical backtests, cache is always valid (history doesn't change)

    # Only refresh if end_date is recent (within 7 days)

    end_dt = datetime.strptime(end_date, '%Y-%m-%d')

    is_historical = (datetime.now() - end_dt).days > 7

    

    if is_historical and use_cache:

        cached = load_from_cache(cache_file)

        if cached and cached.get('prices'):

            return cached['prices'], None

    

    # If no API key, try to use cache

    if not EOD_OK:

        cached = load_from_cache(cache_file)

        if cached and cached.get('prices'):

            log_cache.info(f"Using offline cache for {ticker} prices")

            return cached['prices'], None

        return None, "EOD API not configured and no cached data available"

    

    try:

        eod_ticker = get_eod_ticker(ticker)

        url = f"https://eodhd.com/api/eod/{eod_ticker}?api_token={EOD_API_KEY}&fmt=json&from={start_date}&to={end_date}"

        log_api.info(f"Fetching prices: {ticker} ({start_date} to {end_date})")

        response = requests.get(url, timeout=30)

        

        if response.status_code == 200:

            data = response.json()

            save_to_cache(cache_file, {'prices': data})

            return data, None

        else:

            cached = load_from_cache(cache_file)

            if cached and cached.get('prices'):

                log_cache.info(f"API error, using cached data for {ticker}")

                return cached['prices'], None

            return None, f"API error: {response.status_code}"

            

    except Exception as e:

        cached = load_from_cache(cache_file)

        if cached and cached.get('prices'):

            log_cache.info(f"Exception, using cached data for {ticker}")

            return cached['prices'], None

        return None, str(e)


def eod_get_dividends(ticker: str, start_date: str, end_date: str, use_cache: bool = True) -> Tuple[Optional[List[Dict[str, Any]]], Optional[str]]:
    """Get dividend history from EOD API"""
    cache_key = f"div_{ticker}_{start_date}_{end_date}"
    cache_file = get_cache_path('dividends', cache_key)

    # Try cache first
    if use_cache and is_cache_valid(cache_file, 7):  # 7 days cache for dividends
        cached = load_from_cache(cache_file)
        if cached and cached.get('dividends') is not None:
            return cached['dividends'], None

    if not EOD_OK:
        cached = load_from_cache(cache_file)
        if cached and cached.get('dividends') is not None:
            return cached['dividends'], None
        return None, "EOD API not configured"

    try:
        eod_ticker = get_eod_ticker(ticker)
        url = f"https://eodhd.com/api/div/{eod_ticker}?api_token={EOD_API_KEY}&fmt=json&from={start_date}&to={end_date}"
        log_api.info(f"Fetching dividends: {ticker} ({start_date} to {end_date})")
        response = requests.get(url, timeout=30)

        if response.status_code == 200:
            data = response.json()
            save_to_cache(cache_file, {'dividends': data})
            return data, None
        else:
            cached = load_from_cache(cache_file)
            if cached and cached.get('dividends') is not None:
                return cached['dividends'], None
            return [], None  # Return empty list if no dividends

    except Exception as e:
        cached = load_from_cache(cache_file)
        if cached and cached.get('dividends') is not None:
            return cached['dividends'], None
        return None, str(e)


def download_all_data(scope='france', start_date='2010-01-01', progress_callback=None):

    """

    Pre-download all data for offline use

    Returns dict with success count, error count, and errors list

    """

    result = {

        'success': 0,

        'errors': 0,

        'error_list': [],

        'total': 0

    }

    

    if not EOD_OK:

        result['error_list'].append("EOD API not configured")

        return result

    

    # Get universe

    if scope == 'europe':

        exchanges = ['PA', 'AS', 'BR', 'MI', 'MC', 'XETRA', 'SW', 'LSE']

    else:

        exchanges = ['PA']

    

    all_tickers = []

    for exchange in exchanges:

        tickers = load_universe_from_eod(exchange)

        all_tickers.extend([t['ticker'] for t in tickers])

    

    result['total'] = len(all_tickers)

    end_date = datetime.now().strftime('%Y-%m-%d')

    

    for i, ticker in enumerate(all_tickers):

        if progress_callback:

            progress_callback(i + 1, len(all_tickers), ticker)

        

        # Download fundamentals

        fund, err = eod_get_fundamentals(ticker, use_cache=True)

        if err:

            result['errors'] += 1

            result['error_list'].append(f"{ticker} fundamentals: {err}")

            continue

        

        # Download prices

        prices, err = eod_get_historical_prices(ticker, start_date, end_date, use_cache=True)

        if err:

            result['errors'] += 1

            result['error_list'].append(f"{ticker} prices: {err}")

            continue

        

        result['success'] += 1

    

    return result

    

    # Par defaut, ajouter .PA pour Paris

    return ticker + '.PA'



def eod_get_historical_prices(ticker: str, start_date: str, end_date: str, use_cache: bool = True) -> Tuple[Optional[List[Dict[str, Any]]], Optional[str]]:

    """Get historical prices from EOD"""

    if not EOD_OK:

        return None, "EOD API not configured. Set EOD_API_KEY environment variable."

    

    cache_dir = CONFIG['backtest_cache_dir']

    os.makedirs(cache_dir, exist_ok=True)

    cache_file = os.path.join(cache_dir, f"prices_{ticker.replace('.', '_')}_{start_date}_{end_date}.json")

    

    if use_cache and os.path.exists(cache_file):

        try:

            return json.load(open(cache_file, encoding='utf-8')), None

        except Exception:
            pass

    

    try:

        eod_ticker = get_eod_ticker(ticker)

        url = f"https://eodhd.com/api/eod/{eod_ticker}?api_token={EOD_API_KEY}&fmt=json&from={start_date}&to={end_date}"

        log_api.info(f"Fetching prices: {eod_ticker} from {start_date} to {end_date}")

        response = requests.get(url, timeout=30)

        

        if response.status_code == 200:

            data = response.json()

            if data:  # Only cache if we got data

                with open(cache_file, 'w', encoding='utf-8') as f:

                    json.dump(data, f)

            return data, None

        else:

            return None, f"API error: {response.status_code} - {response.text[:100]}"

            

    except Exception as e:

        return None, str(e)



def extract_historical_fundamentals(fundamentals):

    """Extract historical PE, ROE, margins from EOD fundamentals"""

    history = []

    

    if not fundamentals:

        return history

    

    # Financials historiques

    financials = fundamentals.get('Financials', {})

    

    # Income Statement

    income = financials.get('Income_Statement', {}).get('yearly', {})

    # Balance Sheet

    balance = financials.get('Balance_Sheet', {}).get('yearly', {})

    # Cash Flow Statement

    cashflow = financials.get('Cash_Flow', {}).get('yearly', {})

    

    years = sorted(set(list(income.keys()) + list(balance.keys()) + list(cashflow.keys())), reverse=True)

    

    for year in years:

        try:

            inc = income.get(year, {})

            bal = balance.get(year, {})

            

            # Net Income

            net_income = float(inc.get('netIncome', 0) or 0)

            

            # Revenue

            revenue = float(inc.get('totalRevenue', 0) or 0)

            

            # Total Equity

            equity = float(bal.get('totalStockholderEquity', 0) or 0)

            

            # Total Assets

            assets = float(bal.get('totalAssets', 0) or 0)

            

            # Calculs

            roe = net_income / equity if equity > 0 else None

            roa = net_income / assets if assets > 0 else None

            net_margin = net_income / revenue if revenue > 0 else None

            

            # Operating Income

            op_income = float(inc.get('operatingIncome', 0) or 0)

            op_margin = op_income / revenue if revenue > 0 else None

            

            # Debt

            total_debt = float(bal.get('longTermDebt', 0) or 0) + float(bal.get('shortTermDebt', 0) or 0)

            debt_equity = total_debt / equity if equity > 0 else None

            

            # Current Ratio

            current_assets = float(bal.get('totalCurrentAssets', 0) or 0)

            current_liab = float(bal.get('totalCurrentLiabilities', 0) or 0)

            current_ratio = current_assets / current_liab if current_liab > 0 else None

            

            # EBITDA

            ebitda = float(inc.get('ebitda', 0) or 0)

            

            # Free Cash Flow (approximation)

            fcf = float(inc.get('freeCashFlow', 0) or 0) if 'freeCashFlow' in inc else None



            # Operating Cash Flow from Cash Flow Statement

            cf = cashflow.get(year, {})

            operating_cashflow = float(cf.get('totalCashFromOperatingActivities', 0) or 0)

            if operating_cashflow == 0:

                operating_cashflow = float(cf.get('operatingCashFlow', 0) or 0)



            history.append({

                'year': year[:4] if len(year) >= 4 else year,

                'date': year,

                'revenue': revenue,

                'net_income': net_income,

                'ebitda': ebitda,

                'equity': equity,

                'total_assets': assets,

                'total_debt': total_debt,

                'roe': roe,

                'roa': roa,

                'net_margin': net_margin,

                'op_margin': op_margin,

                'debt_equity': debt_equity,

                'current_ratio': current_ratio,

                'fcf': fcf,

                'operating_cashflow': operating_cashflow

            })

            

        except Exception as e:

            log.error(f"Error parsing year {year}: {e}")

            continue

    

    return history



def get_market_cap_history(ticker, prices, shares_outstanding):

    """Calculate historical market cap from prices and shares"""

    if not prices or not shares_outstanding:

        return []

    

    return [

        {

            'date': p['date'],

            'market_cap': p['close'] * shares_outstanding,

            'price': p['close']

        }

        for p in prices

    ]



# =============================================================================

# BACKTEST ENGINE

# =============================================================================



def run_backtest(params: Dict[str, Any]) -> Dict[str, Any]:

    """

    Run a DYNAMIC Higgons-style value investing backtest

    

    At each rebalancing:

    1. Screen the ENTIRE universe for stocks meeting criteria

    2. Rank by Higgons score

    3. Sell positions that no longer qualify (PE > sell_pe, ROE degraded, etc.)

    4. Buy best new opportunities to maintain target position count

    """

    

    results = {

        'params': params,

        'trades': [],

        'equity_curve': [],

        'positions_history': [],

        'metrics': {},

        'yearly_returns': [],

        'benchmark_curve': [],

        'screening_history': [],

        'errors': []

    }

    

    if not EOD_OK:

        results['errors'].append("EOD API not configured. Set EOD_API_KEY environment variable.")

        return results

    

    # Parse parameters

    start_date = params.get('start_date', '2015-01-01')

    end_date = params.get('end_date', datetime.now().strftime('%Y-%m-%d'))

    universe_scope = params.get('universe_scope', 'france')  # 'france' or 'europe'

    

    # Higgons BUY criteria

    pe_max_buy = float(params.get('pe_max', 12))

    roe_min_buy = float(params.get('roe_min', 10)) / 100

    debt_equity_max = float(params.get('debt_equity_max', 100)) / 100

    pcf_max_buy = float(params.get('pcf_max', 10))  # P/CF maximum for buying



    # Higgons SELL criteria

    pe_sell_threshold = float(params.get('pe_sell', 17))  # Sell if PE > 17

    roe_min_hold = float(params.get('roe_min_hold', 8)) / 100  # Sell if ROE < 8%

    

    rebalance_freq = params.get('rebalance_freq', 'quarterly')

    initial_capital = float(params.get('initial_capital', 100000))

    max_positions = int(params.get('max_positions', 20))

    benchmark = params.get('benchmark', '^FCHI')

    

    # Get universe of tickers to scan

    if universe_scope == 'europe':

        if not UNIVERSE_EUROPE:

            init_universes()

        universe_tickers = [t['ticker'] for t in UNIVERSE_EUROPE]

    else:

        if not UNIVERSE_FRANCE:

            init_universes()

        universe_tickers = [t['ticker'] for t in UNIVERSE_FRANCE]

    

    # If custom universe provided, use that instead

    custom_universe = params.get('universe', [])

    if custom_universe and len(custom_universe) > 0:

        universe_tickers = custom_universe

    

    log_backtest.info(f"Universe: {len(universe_tickers)} tickers ({universe_scope})")

    log_backtest.info(f"Period: {start_date} to {end_date}")

    log_backtest.info(f"Buy criteria: PE <= {pe_max_buy}, ROE >= {roe_min_buy*100}%, P/CF <= {pcf_max_buy}")

    log_backtest.info(f"Sell criteria: PE > {pe_sell_threshold} or ROE < {roe_min_hold*100}%")

    

    # Pre-load fundamental data for all tickers (with caching)

    log_backtest.info("Loading fundamental data...")

    ticker_fundamentals = {}

    loaded_count = 0

    

    for ticker in universe_tickers[:200]:  # Limit to 200 for now to avoid API overload

        fund, err = eod_get_fundamentals(ticker, use_cache=True)

        if fund and not err:

            hist = extract_historical_fundamentals(fund)

            if hist:

                ticker_fundamentals[ticker] = {

                    'name': fund.get('General', {}).get('Name', ticker),

                    'sector': fund.get('General', {}).get('Sector', 'Unknown'),

                    'fundamentals': hist,

                    'fund_by_year': {f['year']: f for f in hist}

                }

                loaded_count += 1

        

        if loaded_count % 20 == 0:

            log_backtest.info(f"Loaded {loaded_count} tickers with valid data...")

    

    log_backtest.info(f"Loaded fundamentals for {loaded_count} tickers")

    

    if loaded_count == 0:

        results['errors'].append("No valid fundamental data found for any ticker")

        return results

    

    # Load price data for tickers with fundamentals

    log_backtest.info("Loading price data...")

    ticker_prices = {}

    

    for ticker in ticker_fundamentals.keys():

        prices, err = eod_get_historical_prices(ticker, start_date, end_date, use_cache=True)

        if prices and len(prices) > 0:

            ticker_prices[ticker] = {p['date']: p for p in prices}

    

    log_backtest.info(f"Loaded prices for {len(ticker_prices)} tickers")

    

    # Load benchmark

    eod_benchmark = get_eod_ticker(benchmark)

    log_backtest.info(f"Loading benchmark {benchmark} -> {eod_benchmark}...")

    bench_prices_raw, err = eod_get_historical_prices(benchmark, start_date, end_date)

    if bench_prices_raw and len(bench_prices_raw) > 0:

        results['benchmark_curve'] = [{'date': p['date'], 'price': p['close']} for p in bench_prices_raw]

        log_backtest.info(f"Benchmark loaded: {len(results['benchmark_curve'])} data points")

    else:

        log_backtest.warning(f"Could not load benchmark {benchmark}: {err}")

        results['errors'].append(f"Benchmark {benchmark} not loaded: {err}")

    

    # Generate rebalancing dates

    rebalance_dates = generate_rebalance_dates(start_date, end_date, rebalance_freq)

    log_backtest.info(f"{len(rebalance_dates)} rebalancing periods")

    

    # Simulation

    capital = initial_capital

    positions = {}  # ticker -> {'shares': n, 'entry_price': p, 'entry_date': d, 'entry_pe': pe, 'entry_roe': roe}

    equity_curve = []

    

    for rebal_idx, rebal_date in enumerate(rebalance_dates):

        year = rebal_date[:4]

        fund_year = str(int(year) - 1)  # Use previous year's fundamentals

        

        log_backtest.info(f"[{rebal_idx+1}/{len(rebalance_dates)}] Rebalancing {rebal_date}...")

        

        # ============================================

        # STEP 1: SCREEN ENTIRE UNIVERSE

        # ============================================

        screened = []

        

        for ticker, data in ticker_fundamentals.items():

            if ticker not in ticker_prices:

                continue

            

            # Get fundamentals for the relevant year

            fund = data['fund_by_year'].get(fund_year) or data['fund_by_year'].get(year)

            if not fund:

                continue

            

            # Get current price

            price = get_price_on_date(ticker_prices[ticker], rebal_date)

            if not price or price <= 0:

                continue

            

            # Extract metrics

            roe = fund.get('roe')

            net_margin = fund.get('net_margin')

            debt_eq = fund.get('debt_equity')

            net_income = fund.get('net_income', 0)

            equity = fund.get('equity', 0)

            operating_cashflow = fund.get('operating_cashflow', 0)

            

            # Calculate PE (approximation using earnings yield)

            # PE = Market Cap / Net Income

            # Since we don't have shares, use: PE â‰ˆ 1/ROE when P/B â‰ˆ 1

            # Better approximation: use net_margin and revenue trends

            if roe and roe > 0:

                pe_approx = 1 / roe  # Simplified PE proxy

            else:

                pe_approx = None



            # Calculate P/CF proxy: Book Value / Operating Cash Flow

            # Lower is better (similar to PE - lower means cheaper on cash flow basis)

            pcf_proxy = None

            if equity and equity > 0 and operating_cashflow and operating_cashflow > 0:

                pcf_proxy = equity / operating_cashflow



            # Check if meets BUY criteria

            if roe is None or roe < roe_min_buy:

                continue

            if pe_approx and pe_approx > pe_max_buy:

                continue

            if debt_eq and debt_eq > debt_equity_max:

                continue

            if pcf_proxy and pcf_proxy > pcf_max_buy:

                continue



            # Calculate Higgons Score (now includes P/CF)

            score = calculate_higgons_score_for_backtest(pe_approx, roe, debt_eq, net_margin, pcf_proxy)



            screened.append({

                'ticker': ticker,

                'name': data['name'],

                'sector': data['sector'],

                'price': price,

                'pe': pe_approx,

                'roe': roe,

                'debt_equity': debt_eq,

                'net_margin': net_margin,

                'pcf': pcf_proxy,

                'score': score

            })

        

        # Sort by score (best first)

        screened.sort(key=lambda x: x['score'], reverse=True)

        

        results['screening_history'].append({

            'date': rebal_date,

            'candidates': len(screened),

            'top_10': [{'ticker': s['ticker'], 'score': s['score'], 'pe': s['pe'], 'roe': s['roe'], 'pcf': s.get('pcf')} for s in screened[:10]]

        })

        

        # ============================================

        # STEP 2: CHECK EXISTING POSITIONS FOR SELL

        # ============================================

        positions_to_sell = []

        

        for ticker, pos in positions.items():

            should_sell = False

            sell_reason = ""

            

            if ticker not in ticker_fundamentals or ticker not in ticker_prices:

                should_sell = True

                sell_reason = "No data"

            else:

                # Get current fundamentals

                fund = ticker_fundamentals[ticker]['fund_by_year'].get(fund_year) or \
                       ticker_fundamentals[ticker]['fund_by_year'].get(year)

                

                if fund:

                    current_roe = fund.get('roe')

                    current_pe = 1 / current_roe if current_roe and current_roe > 0 else None

                    

                    # SELL if PE became too high

                    if current_pe and current_pe > pe_sell_threshold:

                        should_sell = True

                        sell_reason = f"PE too high ({current_pe:.1f} > {pe_sell_threshold})"

                    

                    # SELL if ROE degraded

                    elif current_roe and current_roe < roe_min_hold:

                        should_sell = True

                        sell_reason = f"ROE too low ({current_roe*100:.1f}% < {roe_min_hold*100}%)"

                    

                    # SELL if no longer in top candidates (underperformer)

                    elif ticker not in [s['ticker'] for s in screened[:max_positions * 2]]:

                        should_sell = True

                        sell_reason = "No longer top ranked"

            

            if should_sell:

                positions_to_sell.append((ticker, sell_reason))

        

        # Execute SELLS

        for ticker, reason in positions_to_sell:

            pos = positions[ticker]

            price = get_price_on_date(ticker_prices.get(ticker, {}), rebal_date)

            

            if price:

                sell_value = pos['shares'] * price

                pnl_pct = (price / pos['entry_price'] - 1) * 100

                capital += sell_value

                

                results['trades'].append({

                    'date': rebal_date,

                    'action': 'SELL',

                    'ticker': ticker,

                    'name': ticker_fundamentals.get(ticker, {}).get('name', ticker),

                    'shares': pos['shares'],

                    'price': price,

                    'value': sell_value,

                    'pnl_pct': pnl_pct,

                    'entry_price': pos['entry_price'],

                    'entry_date': pos['entry_date'],

                    'reason': reason

                })

            

            del positions[ticker]

        

        # ============================================

        # STEP 3: BUY NEW POSITIONS

        # ============================================

        positions_to_fill = max_positions - len(positions)

        

        if positions_to_fill > 0 and screened:

            # Filter out already held positions

            available = [s for s in screened if s['ticker'] not in positions]

            

            # Take top N available

            to_buy = available[:positions_to_fill]

            

            if to_buy and capital > 0:

                capital_per_position = capital / len(to_buy)

                

                for candidate in to_buy:

                    ticker = candidate['ticker']

                    price = candidate['price']

                    

                    shares = int(capital_per_position / price)

                    if shares > 0:

                        cost = shares * price

                        capital -= cost

                        

                        positions[ticker] = {

                            'shares': shares,

                            'entry_price': price,

                            'entry_date': rebal_date,

                            'entry_pe': candidate['pe'],

                            'entry_roe': candidate['roe']

                        }

                        

                        results['trades'].append({

                            'date': rebal_date,

                            'action': 'BUY',

                            'ticker': ticker,

                            'name': candidate['name'],

                            'shares': shares,

                            'price': price,

                            'value': cost,

                            'pe': candidate['pe'],

                            'roe': candidate['roe'],

                            'score': candidate['score']

                        })

        

        # ============================================

        # STEP 4: CALCULATE PORTFOLIO VALUE

        # ============================================

        total_value = capital

        pos_details = []

        

        for ticker, pos in positions.items():

            price = get_price_on_date(ticker_prices.get(ticker, {}), rebal_date)

            if price:

                val = pos['shares'] * price

                total_value += val

                pos_details.append({

                    'ticker': ticker,

                    'shares': pos['shares'],

                    'price': price,

                    'value': val,

                    'pnl_pct': (price / pos['entry_price'] - 1) * 100

                })

        

        equity_curve.append({

            'date': rebal_date,

            'value': total_value,

            'cash': capital,

            'invested': total_value - capital,

            'positions': len(positions),

            'sells': len(positions_to_sell),

            'buys': len([t for t in results['trades'] if t['date'] == rebal_date and t['action'] == 'BUY'])

        })

        

        results['positions_history'].append({

            'date': rebal_date,

            'positions': pos_details.copy(),

            'screened_count': len(screened)

        })

    

    results['equity_curve'] = equity_curve

    

    # Calculate final metrics

    if equity_curve:

        results['metrics'] = calculate_backtest_metrics(

            equity_curve, 

            results['benchmark_curve'],

            initial_capital,

            results['trades']

        )

    

    results['yearly_returns'] = calculate_yearly_returns(equity_curve)

    

    return results





def calculate_higgons_score_for_backtest(pe: Optional[float], roe: Optional[float], debt_equity: Optional[float], net_margin: Optional[float], pcf: Optional[float] = None) -> int:

    """Calculate Higgons-style score for ranking (max 115 points with P/CF)"""

    score = 0



    # PE Score (0-35 points) - lower is better

    if pe and pe > 0:

        if pe < 7: score += 35

        elif pe < 10: score += 30

        elif pe < 12: score += 25

        elif pe < 15: score += 15

        elif pe < 20: score += 5



    # P/CF Score (0-15 points) - lower is better

    if pcf and pcf > 0:

        if pcf <= 6: score += 15

        elif pcf <= 8: score += 12

        elif pcf <= 10: score += 9

        elif pcf <= 12: score += 6

        elif pcf <= 15: score += 3



    # ROE Score (0-30 points) - higher is better

    if roe and roe > 0:

        if roe >= 0.25: score += 30

        elif roe >= 0.20: score += 25

        elif roe >= 0.15: score += 20

        elif roe >= 0.12: score += 15

        elif roe >= 0.10: score += 10



    # Debt/Equity Score (0-20 points) - lower is better

    if debt_equity is not None:

        if debt_equity <= 0: score += 20  # Net cash

        elif debt_equity <= 0.3: score += 15

        elif debt_equity <= 0.5: score += 10

        elif debt_equity <= 1.0: score += 5



    # Net Margin Score (0-15 points) - higher is better

    if net_margin and net_margin > 0:

        if net_margin >= 0.15: score += 15

        elif net_margin >= 0.10: score += 10

        elif net_margin >= 0.05: score += 5



    return score





def generate_rebalance_dates(start_date, end_date, freq):

    """Generate list of rebalancing dates"""

    dates = []

    current = datetime.strptime(start_date, '%Y-%m-%d')

    end = datetime.strptime(end_date, '%Y-%m-%d')

    

    if freq == 'monthly':

        delta = 1

    elif freq == 'quarterly':

        delta = 3

    elif freq == 'semi-annual':

        delta = 6

    else:  # yearly

        delta = 12

    

    while current <= end:

        dates.append(current.strftime('%Y-%m-%d'))

        

        # Avancer de N mois

        month = current.month + delta

        year = current.year

        while month > 12:

            month -= 12

            year += 1

        

        try:

            current = current.replace(year=year, month=month)

        except ValueError:

            # Gerer les fins de mois

            current = current.replace(year=year, month=month, day=28)

    

    return dates





def get_price_on_date(prices_dict, target_date):

    """Get price on specific date or closest previous date"""

    if target_date in prices_dict:

        return prices_dict[target_date]['close']

    

    # Chercher la date la plus proche avant

    target = datetime.strptime(target_date, '%Y-%m-%d')

    closest_date = None

    closest_price = None

    

    for date_str, data in prices_dict.items():

        date = datetime.strptime(date_str, '%Y-%m-%d')

        if date <= target:

            if closest_date is None or date > closest_date:

                closest_date = date

                closest_price = data['close']

    

    return closest_price





def calculate_backtest_metrics(equity_curve, benchmark_curve, initial_capital, trades):

    """Calculate performance metrics"""

    metrics = {}

    

    if not equity_curve:

        return metrics

    

    # Total Return

    final_value = equity_curve[-1]['value']

    total_return = (final_value / initial_capital - 1) * 100

    metrics['total_return'] = total_return

    

    # CAGR

    start = datetime.strptime(equity_curve[0]['date'], '%Y-%m-%d')

    end = datetime.strptime(equity_curve[-1]['date'], '%Y-%m-%d')

    years = (end - start).days / 365.25

    

    if years > 0:

        cagr = ((final_value / initial_capital) ** (1 / years) - 1) * 100

        metrics['cagr'] = cagr

    else:

        metrics['cagr'] = 0

    

    # Max Drawdown

    peak = initial_capital

    max_dd = 0

    

    for point in equity_curve:

        value = point['value']

        if value > peak:

            peak = value

        dd = (peak - value) / peak * 100

        if dd > max_dd:

            max_dd = dd

    

    metrics['max_drawdown'] = max_dd

    

    # Volatility (annualized)

    returns = []

    for i in range(1, len(equity_curve)):

        ret = (equity_curve[i]['value'] / equity_curve[i-1]['value']) - 1

        returns.append(ret)

    

    if returns:

        import statistics

        vol = statistics.stdev(returns) * (12 ** 0.5) * 100  # Annualized (assuming monthly)

        metrics['volatility'] = vol

        

        # Sharpe Ratio (assuming 2% risk-free rate)

        rf = 0.02

        if vol > 0:

            sharpe = (metrics['cagr'] / 100 - rf) / (vol / 100)

            metrics['sharpe'] = sharpe

        else:

            metrics['sharpe'] = 0

    else:

        metrics['volatility'] = 0

        metrics['sharpe'] = 0

    

    # Win Rate

    winning_trades = [t for t in trades if t.get('action') == 'SELL' and t.get('pnl_pct', 0) > 0]

    losing_trades = [t for t in trades if t.get('action') == 'SELL' and t.get('pnl_pct', 0) <= 0]

    total_closed = len(winning_trades) + len(losing_trades)

    

    if total_closed > 0:

        metrics['win_rate'] = len(winning_trades) / total_closed * 100

        metrics['total_trades'] = total_closed

        

        # Average win/loss

        if winning_trades:

            metrics['avg_win'] = sum(t['pnl_pct'] for t in winning_trades) / len(winning_trades)

        else:

            metrics['avg_win'] = 0

            

        if losing_trades:

            metrics['avg_loss'] = sum(t['pnl_pct'] for t in losing_trades) / len(losing_trades)

        else:

            metrics['avg_loss'] = 0

    else:

        metrics['win_rate'] = 0

        metrics['total_trades'] = 0

        metrics['avg_win'] = 0

        metrics['avg_loss'] = 0

    

    # Benchmark comparison

    if benchmark_curve and len(benchmark_curve) >= 2:

        bench_start = benchmark_curve[0]['price']

        bench_end = benchmark_curve[-1]['price']

        bench_return = (bench_end / bench_start - 1) * 100

        metrics['benchmark_return'] = bench_return

        metrics['alpha'] = total_return - bench_return

    else:

        metrics['benchmark_return'] = 0

        metrics['alpha'] = total_return

    

    metrics['final_value'] = final_value

    metrics['initial_capital'] = initial_capital

    

    return metrics





def calculate_yearly_returns(equity_curve):

    """Calculate returns by year"""

    yearly = {}

    

    for point in equity_curve:

        year = point['date'][:4]

        if year not in yearly:

            yearly[year] = {'start': point['value'], 'end': point['value']}

        else:

            yearly[year]['end'] = point['value']

    

    returns = []

    prev_end = None

    

    for year in sorted(yearly.keys()):

        data = yearly[year]

        start = prev_end if prev_end else data['start']

        end = data['end']

        ret = (end / start - 1) * 100 if start > 0 else 0

        returns.append({'year': year, 'return': ret, 'end_value': end})

        prev_end = end

    

    return returns



def create_memo_docx(ticker, name, sector, country, signal, target_price, thesis, strengths, risks, valuation, notes):

    """Create a new investment memo as a .docx file"""

    try:

        from docx import Document

        from docx.shared import Inches, Pt

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        

        doc = Document()

        

        # Title

        title = doc.add_heading(f'Investment Memo: {name}', 0)

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        

        # Subtitle

        subtitle = doc.add_paragraph(f'{ticker} | {sector} | {country}')

        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        

        # Signal

        doc.add_heading('SIGNAL', level=1)

        signal_para = doc.add_paragraph()

        signal_para.add_run(signal.upper()).bold = True

        if target_price:

            signal_para.add_run(f'  |  Target Price: {target_price} EUR')

        

        # Date

        doc.add_paragraph(f'Date: {datetime.now().strftime("%d %B %Y")}')

        doc.add_paragraph()

        

        # Thesis

        if thesis:

            doc.add_heading('INVESTMENT THESIS', level=1)

            doc.add_paragraph(thesis)

        

        # Strengths

        if strengths:

            doc.add_heading('KEY STRENGTHS / CATALYSTS', level=1)

            for line in strengths.split('\n'):

                if line.strip():

                    doc.add_paragraph(line.strip(), style='List Bullet')

        

        # Risks

        if risks:

            doc.add_heading('RISKS / CONCERNS', level=1)

            for line in risks.split('\n'):

                if line.strip():

                    doc.add_paragraph(line.strip(), style='List Bullet')

        

        # Valuation

        if valuation:

            doc.add_heading('VALUATION', level=1)

            doc.add_paragraph(valuation)

        

        # Notes

        if notes:

            doc.add_heading('ADDITIONAL NOTES', level=1)

            doc.add_paragraph(notes)

        

        # Save

        filename = f'Memo_{ticker.replace(".", "_")}.docx'

        filepath = os.path.join(CONFIG['memo_dir'], filename)

        doc.save(filepath)

        return filepath, None

        

    except ImportError:

        # python-docx not installed, create a simple text file instead

        try:

            filename = f'Memo_{ticker.replace(".", "_")}.txt'

            filepath = os.path.join(CONFIG['memo_dir'], filename)

            

            content = f"""INVESTMENT MEMO: {name}

{'='*50}

{ticker} | {sector} | {country}

Date: {datetime.now().strftime("%d %B %Y")}



SIGNAL: {signal.upper()}

Target Price: {target_price or 'N/A'} EUR



{'='*50}

INVESTMENT THESIS

{'='*50}

{thesis or 'N/A'}



{'='*50}

KEY STRENGTHS / CATALYSTS

{'='*50}

{strengths or 'N/A'}



{'='*50}

RISKS / CONCERNS

{'='*50}

{risks or 'N/A'}



{'='*50}

VALUATION

{'='*50}

{valuation or 'N/A'}



{'='*50}

ADDITIONAL NOTES

{'='*50}

{notes or 'N/A'}

"""

            with open(filepath, 'w', encoding='utf-8') as f:

                f.write(content)

            return filepath, None

            

        except Exception as e:

            return None, str(e)

    except Exception as e:

        return None, str(e)





def generate_memo_with_ai(security_data):

    """Generate investment memo using Claude AI"""

    if not ANTHROPIC_OK:

        return None, "Anthropic API not configured. Set ANTHROPIC_API_KEY environment variable."

    

    try:

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        

        # Build the prompt with all available data

        prompt = f"""Tu es un analyste financier senior specialise dans les small et mid caps europeennes. 

Analyse les donnees suivantes et genere un memo d'investissement professionnel et structure.



=== DONNEES DE LA SOCIETE ===

Nom: {security_data.get('name', 'N/A')}

Ticker: {security_data.get('ticker', 'N/A')}

Secteur: {security_data.get('sector', 'N/A')}

Industrie: {security_data.get('industry', 'N/A')}

Pays: {security_data.get('country', 'N/A')}



=== DONNEES DE PRIX ===

Prix actuel: {security_data.get('price', 'N/A')} EUR

Variation jour: {security_data.get('change_pct', 'N/A')}%

Plus haut 52 semaines: {security_data.get('high_52w', 'N/A')} EUR

Plus bas 52 semaines: {security_data.get('low_52w', 'N/A')} EUR

Target Price analystes: {security_data.get('target_price', 'N/A')} EUR



=== VALORISATION ===

P/E (TTM): {security_data.get('pe', 'N/A')}

P/E Forward: {security_data.get('forward_pe', 'N/A')}

EPS: {security_data.get('eps', 'N/A')}

Book Value: {security_data.get('book_value', 'N/A')}



=== RENTABILITE ===

ROE: {security_data.get('roe', 'N/A')}

Marge operationnelle: {security_data.get('profit_margin', 'N/A')}

Rendement dividende: {security_data.get('dividend_yield', 'N/A')}



=== BILAN ===

Market Cap: {security_data.get('market_cap', 'N/A')}

Revenue: {security_data.get('revenue', 'N/A')}

Debt/Equity: {security_data.get('debt_equity', 'N/A')}

Current Ratio: {security_data.get('current_ratio', 'N/A')}



=== DESCRIPTION ===

{security_data.get('description', 'N/A')[:1500]}



=== INSTRUCTIONS ===

Genere un memo d'investissement avec EXACTEMENT cette structure (utilise ces titres):



SIGNAL

[Donne une recommandation: ACHAT, SURVEILLANCE, NEUTRE, ou ECARTER avec un prix cible si possible]



RESUME EXECUTIF

[2-3 phrases resumant la these d'investissement]



THÃˆSE D'INVESTISSEMENT

[Paragraphe detaille expliquant pourquoi investir ou non]



POINTS FORTS

[Liste a puces des 4-5 principaux atouts]



RISQUES

[Liste a puces des 4-5 principaux risques]



ANALYSE DE VALORISATION

[Analyse du PE, ROE, comparaison avec le secteur, est-ce cher ou pas ?]



CATALYSEURS POTENTIELS

[Liste des evenements qui pourraient faire bouger le cours]



CONCLUSION

[Synthese finale avec la recommandation]



Sois precis, factuel et utilise les donnees fournies. Si une donnee est manquante (N/A), mentionne-le comme une limitation.

Ecris en francais."""



        message = client.messages.create(

            model="claude-sonnet-4-20250514",

            max_tokens=4096,

            messages=[

                {"role": "user", "content": prompt}

            ]

        )

        

        ai_response = message.content[0].text

        

        # Parse the AI response to extract sections

        sections = parse_ai_memo_response(ai_response)

        

        # Create the docx file

        filepath, error = create_ai_memo_docx(security_data, sections, ai_response)

        

        return filepath, error

        

    except Exception as e:

        return None, str(e)





def parse_ai_memo_response(response):

    """Parse AI response to extract sections"""

    sections = {

        'signal': '',

        'summary': '',

        'thesis': '',

        'strengths': '',

        'risks': '',

        'valuation': '',

        'catalysts': '',

        'conclusion': ''

    }

    

    # Simple parsing based on section headers

    current_section = None

    current_content = []

    

    for line in response.split('\n'):

        line_upper = line.upper().strip()

        

        if 'SIGNAL' in line_upper and len(line_upper) < 20:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'signal'

            current_content = []

        elif 'RESUME' in line_upper or 'RESUME' in line_upper or 'EXECUTIF' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'summary'

            current_content = []

        elif 'THÃˆSE' in line_upper or 'THESE' in line_upper or 'INVESTISSEMENT' in line_upper and 'THÃˆSE' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'thesis'

            current_content = []

        elif 'POINTS FORTS' in line_upper or 'FORCES' in line_upper or 'ATOUTS' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'strengths'

            current_content = []

        elif 'RISQUES' in line_upper or 'RISKS' in line_upper or 'FAIBLESSES' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'risks'

            current_content = []

        elif 'VALORISATION' in line_upper or 'VALUATION' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'valuation'

            current_content = []

        elif 'CATALYSEUR' in line_upper or 'CATALYST' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'catalysts'

            current_content = []

        elif 'CONCLUSION' in line_upper:

            if current_section:

                sections[current_section] = '\n'.join(current_content).strip()

            current_section = 'conclusion'

            current_content = []

        else:

            if current_section:

                current_content.append(line)

    

    # Don't forget the last section

    if current_section:

        sections[current_section] = '\n'.join(current_content).strip()

    

    return sections





def create_ai_memo_docx(security_data, sections, full_response):

    """Create a docx file from AI-generated memo"""

    try:

        from docx import Document

        from docx.shared import Pt, RGBColor

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        

        doc = Document()

        

        # Title

        title = doc.add_heading(f'Investment Memo: {security_data.get("name", "N/A")}', 0)

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        

        # Subtitle

        subtitle = doc.add_paragraph(f'{security_data.get("ticker", "")} | {security_data.get("sector", "")} | {security_data.get("country", "")}')

        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        

        # Date and AI notice

        doc.add_paragraph(f'Date: {datetime.now().strftime("%d %B %Y")}')

        ai_notice = doc.add_paragraph('📊 Memo generated with AI assistance (Claude)')

        ai_notice.runs[0].font.size = Pt(9)

        ai_notice.runs[0].font.italic = True

        doc.add_paragraph()

        

        # Signal section

        if sections.get('signal'):

            doc.add_heading('SIGNAL', level=1)

            signal_para = doc.add_paragraph()

            signal_text = sections['signal'].strip()

            run = signal_para.add_run(signal_text)

            run.bold = True

        

        # Summary

        if sections.get('summary'):

            doc.add_heading('RESUME EXECUTIF', level=1)

            doc.add_paragraph(sections['summary'])

        

        # Thesis

        if sections.get('thesis'):

            doc.add_heading('THÃˆSE D\'INVESTISSEMENT', level=1)

            doc.add_paragraph(sections['thesis'])

        

        # Strengths

        if sections.get('strengths'):

            doc.add_heading('POINTS FORTS', level=1)

            for line in sections['strengths'].split('\n'):

                line = line.strip()

                if line and not line.startswith('#'):

                    # Remove bullet points if already present

                    if line.startswith(('-', ' - ', '*')):

                        line = line[1:].strip()

                    if line:

                        doc.add_paragraph(line, style='List Bullet')

        

        # Risks

        if sections.get('risks'):

            doc.add_heading('RISQUES', level=1)

            for line in sections['risks'].split('\n'):

                line = line.strip()

                if line and not line.startswith('#'):

                    if line.startswith(('-', ' - ', '*')):

                        line = line[1:].strip()

                    if line:

                        doc.add_paragraph(line, style='List Bullet')

        

        # Valuation

        if sections.get('valuation'):

            doc.add_heading('ANALYSE DE VALORISATION', level=1)

            doc.add_paragraph(sections['valuation'])

        

        # Catalysts

        if sections.get('catalysts'):

            doc.add_heading('CATALYSEURS POTENTIELS', level=1)

            for line in sections['catalysts'].split('\n'):

                line = line.strip()

                if line and not line.startswith('#'):

                    if line.startswith(('-', ' - ', '*')):

                        line = line[1:].strip()

                    if line:

                        doc.add_paragraph(line, style='List Bullet')

        

        # Conclusion

        if sections.get('conclusion'):

            doc.add_heading('CONCLUSION', level=1)

            doc.add_paragraph(sections['conclusion'])

        

        # Key metrics table

        doc.add_heading('DONNEES CLES', level=1)

        table = doc.add_table(rows=6, cols=2)

        table.style = 'Table Grid'

        

        metrics = [

            ('Prix actuel', f"{security_data.get('price', 'N/A')} EUR"),

            ('P/E (TTM)', f"{security_data.get('pe', 'N/A')}"),

            ('ROE', f"{security_data.get('roe', 'N/A')}"),

            ('Market Cap', f"{security_data.get('market_cap', 'N/A')}"),

            ('Rendement div.', f"{security_data.get('dividend_yield', 'N/A')}"),

            ('52W Range', f"{security_data.get('low_52w', 'N/A')} - {security_data.get('high_52w', 'N/A')} EUR"),

        ]

        

        for i, (label, value) in enumerate(metrics):

            table.rows[i].cells[0].text = label

            table.rows[i].cells[1].text = str(value) if value else 'N/A'

        

        # Save

        ticker = security_data.get('ticker', 'UNKNOWN').replace('.', '_')

        filename = f'Memo_{ticker}_AI.docx'

        filepath = os.path.join(CONFIG['memo_dir'], filename)

        doc.save(filepath)

        

        return filepath, None

        

    except ImportError:

        # Fallback to text file

        try:

            ticker = security_data.get('ticker', 'UNKNOWN').replace('.', '_')

            filename = f'Memo_{ticker}_AI.txt'

            filepath = os.path.join(CONFIG['memo_dir'], filename)

            

            with open(filepath, 'w', encoding='utf-8') as f:

                f.write(f"INVESTMENT MEMO: {security_data.get('name', 'N/A')}\n")

                f.write(f"{'='*60}\n")

                f.write(f"Generated with AI (Claude) - {datetime.now().strftime('%d %B %Y')}\n\n")

                f.write(full_response)

            

            return filepath, None

        except Exception as e:

            return None, str(e)

    except Exception as e:

        return None, str(e)



def load_nav_history():

    """Load NAV history from JSON file"""

    return load_json(CONFIG['nav_history_file'], [])



def save_nav_history(history):

    """Save NAV history to JSON file"""

    save_json(CONFIG['nav_history_file'], history)



def update_nav_history(nav, total_cost, pnl, pnl_pct, realized_pnl=0.0):

    """Add today's NAV to history (one entry per day)



    NAV = market value of active positions + cumulative realized PnL

    This ensures that selling positions with gains doesn't show as a NAV drop.

    """

    history = load_nav_history()

    today = datetime.now().strftime('%Y-%m-%d')



    # Check if we already have an entry for today

    existing_idx = None

    for i, entry in enumerate(history):

        if entry['date'] == today:

            existing_idx = i

            break



    new_entry = {

        'date': today,

        'nav': round(nav, 2),

        'cost': round(total_cost, 2),

        'pnl': round(pnl, 2),

        'pnl_pct': round(pnl_pct, 2),

        'realized_pnl': round(realized_pnl, 2)

    }

    

    if existing_idx is not None:

        # Update today's entry

        history[existing_idx] = new_entry

    else:

        # Add new entry

        history.append(new_entry)

    

    # Keep only last 365 days

    history = history[-365:]

    

    save_nav_history(history)

    return history



def add_to_watchlist(ticker, name, country, sector):

    w = load_watchlist()

    if not any(x['ticker'] == ticker for x in w):

        w.append({'ticker': ticker, 'name': name, 'country': country, 'sector': sector, 'added': datetime.now().strftime('%Y-%m-%d')})

        save_watchlist(w)



def remove_from_watchlist(ticker):

    w = [x for x in load_watchlist() if x['ticker'] != ticker]

    save_watchlist(w)



def load_portfolio() -> Tuple[Optional[Any], Optional[str]]:

    if not PANDAS_OK:

        return None, "pandas manquant"

    try:

        df = pd.read_excel(CONFIG['portfolio_file'])

        df.columns = [c.lower().strip().replace(' ', '_') for c in df.columns]

        return df, None

    except Exception as e:

        return None, str(e)


def build_advisor_portfolio_payload(df: Any, cash: float = 0.0, currency: str = 'EUR') -> Dict[str, Any]:

    """Convert current portfolio dataframe to advisor JSON payload."""

    positions: List[Dict[str, Any]] = []

    for _, row in df.iterrows():

        ticker = str(row.get('ticker', '') or '').strip().upper()
        if not ticker:
            continue

        qty = safe_float(row.get('qty', row.get('shares', row.get('quantity', 0)))) or 0.0
        if qty <= 0:
            continue

        avg_price = (
            safe_float(row.get('avg_cost_eur', row.get('avg_price', row.get('average_price'))))
            or 0.0
        )
        current_price = safe_float(row.get('price_eur', row.get('current_price')))

        category = str(row.get('category', row.get('style', 'mixte')) or 'mixte').strip().lower()
        if category not in ['value', 'croissance', 'cyclique', 'mixte']:
            category = 'mixte'

        positions.append({
            'ticker': ticker,
            'name': str(row.get('name', '') or '').strip() or None,
            'isin': str(row.get('isin', '') or '').strip() or None,
            'category': category,
            'shares': qty,
            'avg_price': avg_price,
            'current_price': current_price,
            'sector': str(row.get('sector', '') or '').strip() or None,
            'country': str(row.get('country', '') or '').strip() or None,
            'notes': str(row.get('notes', row.get('thesis', '')) or '').strip() or None,
        })

    return {
        'portfolio': positions,
        'cash': float(cash or 0.0),
        'currency': str(currency or 'EUR').upper(),
    }


def save_portfolio(df):
    """Save portfolio DataFrame to Excel"""
    try:
        df.to_excel(CONFIG['portfolio_file'], index=False)
        return True, None
    except Exception as e:
        return False, str(e)


def add_portfolio_position(ticker, name, qty, avg_cost):
    """Add a new position to the portfolio"""
    df, err = load_portfolio()
    if err:
        return False, err

    # Check if ticker already exists
    if ticker.upper() in df['ticker'].str.upper().values:
        return False, f"{ticker} already exists in portfolio"

    # Create new row
    new_row = {
        'ticker': ticker.upper(),
        'name': name,
        'qty': float(qty),
        'avg_cost_eur': float(avg_cost),
        'price_eur': 0,
        'sector': '',
        'country': ''
    }

    # Add to DataFrame
    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)

    # Save
    success, err = save_portfolio(df)
    return success, err


def edit_portfolio_position(ticker, qty, avg_cost):
    """Edit an existing position in the portfolio"""
    df, err = load_portfolio()
    if err:
        return False, err

    # Find the position
    mask = df['ticker'].str.upper() == ticker.upper()
    if not mask.any():
        return False, f"{ticker} not found in portfolio"

    # Update values
    df.loc[mask, 'qty'] = float(qty)
    df.loc[mask, 'avg_cost_eur'] = float(avg_cost)

    # Save
    success, err = save_portfolio(df)
    return success, err


def remove_portfolio_position(ticker):
    """Remove a position from the portfolio"""
    df, err = load_portfolio()
    if err:
        return False, err

    # Remove the position
    original_len = len(df)
    df = df[df['ticker'].str.upper() != ticker.upper()]

    if len(df) == original_len:
        return False, f"{ticker} not found in portfolio"

    # Save
    success, err = save_portfolio(df)
    return success, err


def safe_float(val):

    """Safely convert to float, return None if invalid"""

    if val is None:

        return None

    if isinstance(val, float) and math.isnan(val):

        return None

    try:

        return float(val)

    except Exception:
        return None



def update_portfolio(df):

    """Update portfolio with ALL yfinance data for Higgons scoring"""

    if not YFINANCE_OK:

        return df

    

    for i, row in df.iterrows():

        try:

            ticker = row['ticker']

            yf_ticker = get_yf_ticker(ticker)

            log_portfolio.info(f"Fetching {yf_ticker}...")

            

            t = yf.Ticker(yf_ticker)

            info = t.info

            

            # === BASIC INFO ===

            # Name

            if pd.isna(row.get('name')) or not str(row.get('name', '')).strip():

                name = info.get('shortName') or info.get('longName')

                if name: df.at[i, 'name'] = name

            

            # Sector & Country

            if pd.isna(row.get('sector')) or not str(row.get('sector', '')).strip():

                sector = info.get('sector')

                if sector: df.at[i, 'sector'] = sector

            

            if pd.isna(row.get('country')) or not str(row.get('country', '')).strip():

                country = info.get('country')

                if country: df.at[i, 'country'] = country

            

            # === PRICE DATA ===

            new_price = info.get('currentPrice') or info.get('regularMarketPrice')

            if new_price and new_price > 0:

                df.at[i, 'price_eur'] = new_price

            

            # === VALUATION METRICS ===

            pe = info.get('trailingPE')

            if pe and pe > 0: df.at[i, 'pe_ttm'] = pe

            

            # === PROFITABILITY ===

            roe = info.get('returnOnEquity')

            if roe is not None: df.at[i, 'roe_ttm'] = roe

            

            roa = info.get('returnOnAssets')

            if roa is not None: df.at[i, 'roa_ttm'] = roa

            

            gross_margin = info.get('grossMargins')

            if gross_margin is not None: df.at[i, 'gross_margin'] = gross_margin

            

            op_margin = info.get('operatingMargins')

            if op_margin is not None: df.at[i, 'operating_margin'] = op_margin

            

            # === FINANCIAL DATA ===

            revenue = info.get('totalRevenue')

            if revenue: df.at[i, 'revenue'] = revenue / 1e6  # Convert to millions

            

            ebitda = info.get('ebitda')

            if ebitda: df.at[i, 'ebitda'] = ebitda

            

            net_income = info.get('netIncomeToCommon')

            if net_income: df.at[i, 'net_income'] = net_income

            

            # === CASH FLOW ===

            op_cashflow = info.get('operatingCashflow')

            if op_cashflow: df.at[i, 'operating_cashflow'] = op_cashflow

            

            free_cashflow = info.get('freeCashflow')

            if free_cashflow: df.at[i, 'free_cashflow'] = free_cashflow

            

            # FCF Yield = Free Cash Flow / Market Cap

            market_cap = info.get('marketCap')

            if free_cashflow and market_cap and market_cap > 0:

                df.at[i, 'fcf_yield'] = free_cashflow / market_cap


            # P/CF = Market Cap / Operating Cash Flow (Price to Cash Flow ratio)
            # Lower is better - indicates company generates cash relative to its valuation
            if op_cashflow and market_cap and op_cashflow > 0:
                pcf = market_cap / op_cashflow
                df.at[i, 'pcf'] = pcf
                df.at[i, 'price_to_cashflow'] = pcf  # Alias for compatibility

            # FCF to Net Income ratio

            if free_cashflow and net_income and net_income != 0:

                df.at[i, 'fcf_to_net_income'] = free_cashflow / abs(net_income)

            

            # === BALANCE SHEET ===

            total_debt = info.get('totalDebt')

            if total_debt is not None: df.at[i, 'total_debt'] = total_debt

            

            total_cash = info.get('totalCash')

            if total_cash is not None: df.at[i, 'total_cash'] = total_cash

            

            total_assets = info.get('totalAssets')

            if total_assets: df.at[i, 'total_assets'] = total_assets

            

            current_ratio = info.get('currentRatio')

            if current_ratio: df.at[i, 'current_ratio'] = current_ratio

            

            # Debt to Equity

            debt_equity = info.get('debtToEquity')

            if debt_equity is not None: df.at[i, 'debt_to_equity'] = debt_equity / 100 if debt_equity > 10 else debt_equity

            

            # Equity Ratio = Total Equity / Total Assets

            total_equity = info.get('totalStockholderEquity')

            if total_equity and total_assets and total_assets > 0:

                df.at[i, 'equity_ratio'] = total_equity / total_assets

            

            # Net Debt to EBITDA

            if total_debt is not None and total_cash is not None and ebitda and ebitda > 0:

                net_debt = total_debt - total_cash

                df.at[i, 'net_debt_to_ebitda'] = net_debt / ebitda

            

            # === SHARES ===

            shares = info.get('sharesOutstanding')

            if shares: df.at[i, 'shares_outstanding'] = shares

            

            # === MOMENTUM (52 week) ===

            high_52w = info.get('fiftyTwoWeekHigh')

            low_52w = info.get('fiftyTwoWeekLow')

            if new_price and high_52w and low_52w:

                # Momentum = where price is in 52w range (-1 to +1)

                if high_52w != low_52w:

                    momentum = (new_price - low_52w) / (high_52w - low_52w) * 2 - 1

                    df.at[i, 'momentum_12m'] = momentum

            

        except Exception as e:

            log_portfolio.error(f"Error {ticker}: {e}")

    

    # === CALCULATE HIGGONS SCORE ===

    df = calc_higgons_score(df)

    

    df.to_excel(CONFIG['portfolio_file'], index=False)

    log_portfolio.info("Portfolio saved!")

    return df





def calc_higgons_score(df):

    """Calculate Higgons score and verdict for each position - Updated logic"""

    for i, row in df.iterrows():

        try:

            pe = safe_float(row.get('pe_ttm')) or 0

            roe = safe_float(row.get('roe_ttm')) or 0

            debt_equity = safe_float(row.get('debt_to_equity')) or 0

            net_debt_ebitda = safe_float(row.get('net_debt_to_ebitda')) or 0

            fcf_yield = safe_float(row.get('fcf_yield')) or 0

            equity_ratio = safe_float(row.get('equity_ratio')) or 0

            momentum = safe_float(row.get('momentum_12m')) or 0

            

            # === HIGGONS GATES ===

            # 1. Profitable Required: PE > 0 and ROE > 0

            profitable = 1 if (pe > 0 and roe > 0) else 0

            df.at[i, 'higgins_profitable_required'] = profitable

            

            # 2. Quality Gate: ROE >= 10%

            quality = 1 if roe >= 0.10 else 0

            df.at[i, 'higgins_quality_gate_pass'] = quality

            

            # 3. Value Gate: PE <= 15 (assoupli)

            value = 1 if (pe > 0 and pe <= 15) else 0

            df.at[i, 'higgins_value_gate_pass'] = value

            

            # 4. Leverage Gate: Net Debt/EBITDA <= 3 or Equity Ratio >= 30%

            leverage = 1 if (net_debt_ebitda <= 3 or equity_ratio >= 0.30) else 0

            df.at[i, 'higgins_leverage_gate_pass'] = leverage

            

            # === HIGGONS SCORE (0-100) ===

            score = 0

            

            # PE Score (0-35 points) - UPDATED

            if pe > 0:

                if pe < 7: score += 35      # Parfait

                elif pe < 10: score += 30   # Top

                elif pe < 12: score += 25   # Bon

                elif pe < 15: score += 15   # Acceptable

                elif pe < 20: score += 5    # Cher

                # > 20 = 0 points (Tres cher)

            

            # ROE Score (0-30 points)

            if roe > 0:

                if roe >= 0.25: score += 30

                elif roe >= 0.20: score += 25

                elif roe >= 0.15: score += 20

                elif roe >= 0.12: score += 15

                elif roe >= 0.10: score += 10

                elif roe >= 0.05: score += 5

            

            # Leverage Score (0-20 points)

            if net_debt_ebitda <= 0: score += 20  # Net cash

            elif net_debt_ebitda <= 1: score += 15

            elif net_debt_ebitda <= 2: score += 10

            elif net_debt_ebitda <= 3: score += 5

            

            # FCF Yield Score (0-10 points)

            if fcf_yield >= 0.10: score += 10

            elif fcf_yield >= 0.05: score += 5

            

            # Equity Ratio Score (0-5 points)

            if equity_ratio >= 0.50: score += 5

            elif equity_ratio >= 0.30: score += 3

            

            df.at[i, 'score_higgins_wo_mom'] = score / 100

            

            # Add momentum adjustment (-10 to +10 points)

            mom_adj = momentum * 10  # -10 to +10

            final_score = (score + mom_adj) / 100

            df.at[i, 'score_higgins'] = max(0, min(1, final_score))

            

            # === CLASSIFY - UPDATED LOGIC ===

            

            # Deep Value Higgons: PE <= 12 and ROE >= 10% and all gates pass

            is_deep_value = 'Oui' if (pe > 0 and pe <= 12 and roe >= 0.10 and profitable and quality and leverage) else 'Non'

            df.at[i, 'is_deep_value_higgins'] = is_deep_value

            

            # Compounder on Sale: ROE >= 15% and PE <= 15

            is_compounder = 'Oui' if (roe >= 0.15 and pe > 0 and pe <= 15) else 'Non'

            df.at[i, 'is_compounder_on_sale'] = is_compounder

            

            # Quality Value (new): PE <= 15 and ROE >= 12%

            is_quality_value = (pe > 0 and pe <= 15 and roe >= 0.12)

            

            # === VERDICT - ALWAYS RECALCULATE ===

            # ACHAT: Deep Value OR Compounder OR (PE < 10 + ROE > 12%)

            if is_deep_value == 'Oui' or is_compounder == 'Oui' or (pe > 0 and pe < 10 and roe >= 0.12):

                verdict = 'Achat'

            

            # SURVEILLANCE: PE <= 15 + ROE >= 10% + profitable

            elif profitable and quality and value:

                verdict = 'Surveillance'

            

            # NEUTRE: Score > 40% but doesn't pass criteria

            elif final_score > 0.40:

                verdict = 'Neutre'

            

            # ECARTER: PE > 20 OR ROE < 5% OR score < 30%

            elif (pe > 20) or (roe > 0 and roe < 0.05) or final_score < 0.30:

                verdict = 'Ecarter'

            

            else:

                verdict = 'Neutre'

            

            df.at[i, 'verdict'] = verdict

            

        except Exception as e:

            log_portfolio.error(f"Error calculating Higgons for row {i}: {e}")

    

    return df



def calc_scores(df):

    """Calculate scores - but use existing 'verdict' column for signal"""

    for i, row in df.iterrows():

        try:

            pe = safe_float(row.get('pe_ttm')) or 0

            roe = safe_float(row.get('roe_ttm')) or 0

            

            # Calculate score for sorting/ranking

            sc = 50

            if pe > 0:

                if pe < 8: sc += 30

                elif pe < 12: sc += 20

                elif pe < 15: sc += 10

                elif pe > 25: sc -= 20

            if roe > 0:

                if roe > 0.20: sc += 20

                elif roe > 0.15: sc += 15

                elif roe > 0.10: sc += 10

            df.at[i, 'score'] = sc

            

            # Use existing verdict column if available, otherwise calculate

            verdict = row.get('verdict', '')

            if verdict and str(verdict).strip():

                df.at[i, 'signal'] = str(verdict).strip()

            else:

                # Fallback calculation only if no verdict

                df.at[i, 'signal'] = 'BUY' if (pe > 0 and pe <= 10 and roe >= 0.10) else ('HOLD' if pe <= 15 else 'SELL')

        except Exception:
            df.at[i, 'score'] = 50

            df.at[i, 'signal'] = row.get('verdict', 'HOLD')

    return df



def load_cache(cache_key='default'):

    """Load cache for a specific key (e.g., 'france_standard', 'europe_standard')"""

    cache_file = CONFIG['screener_cache_file'].replace('.json', f'_{cache_key}.json')

    c = load_json(cache_file, {'date': None, 'data': []})

    if c.get('date'):

        try:

            if (datetime.now() - datetime.strptime(c['date'], '%Y-%m-%d')).days < CONFIG['cache_days']:

                return c['data']

        except Exception:
            pass

    return None



def save_cache(data, cache_key='default'):

    """Save cache for a specific key"""

    cache_file = CONFIG['screener_cache_file'].replace('.json', f'_{cache_key}.json')

    save_json(cache_file, {'date': datetime.now().strftime('%Y-%m-%d'), 'data': data})



def run_screener(force=False, scope='france', use_eod=True, mode='standard'):

    """

    Run the screener on the full universe

    

    Priority:

    1. EOD Historical Data (if API configured or cached data available)

    2. Yahoo Finance (fallback)

    

    scope: 'france', 'europe', or 'legacy' (original 52 stocks)

    mode: 'standard' or 'ai_optimal' (strict criteria from AI optimization)

    """

    

    # AI Optimal criteria (from optimization results)

    AI_OPTIMAL_CRITERIA = {

        'pe_max': 8,

        'roe_min': 12,  # percentage

        'debt_equity_max': 50,  # percentage

        'max_positions': 18

    }

    

    # Check cache first (unless force refresh)

    cache_key = f"{scope}_{mode}"

    if not force:

        c = load_cache(cache_key)

        if c:

            log_screener.info(f"Using cached data for {cache_key}")

            # If AI optimal mode, filter cached results

            if mode == 'ai_optimal':

                return filter_ai_optimal(c, AI_OPTIMAL_CRITERIA)

            return c



    results = []



    # Determine which data source to use

    use_eod_data = use_eod and (EOD_OK or has_eod_cache())



    if use_eod_data and scope in ['france', 'europe']:

        # Use EOD data - scan full universe

        log_screener.info(f"Using EOD data for {scope} universe...")

        results = run_screener_eod(scope)

    else:

        # Fallback to Yahoo Finance with legacy database

        log_screener.info("Using Yahoo Finance with legacy database...")

        results = run_screener_yahoo()



    if results:

        save_cache(results, cache_key)

        log_screener.info(f"Saved {len(results)} results to cache {cache_key}")

    

    # Apply AI optimal filter if requested

    if mode == 'ai_optimal':

        results = filter_ai_optimal(results, AI_OPTIMAL_CRITERIA)

    

    return results





def filter_ai_optimal(results, criteria):

    """Filter results according to AI optimal criteria"""

    filtered = []

    

    for r in results:

        pe = r.get('pe')

        roe = r.get('roe')

        debt_eq = r.get('debt_equity')

        

        # Convert ROE to percentage if needed

        if roe and roe < 1:

            roe_pct = roe * 100

        else:

            roe_pct = roe or 0

        

        # Convert debt_equity to percentage if needed

        if debt_eq and debt_eq < 5:

            debt_pct = debt_eq * 100

        else:

            debt_pct = debt_eq or 0

        

        # Apply strict AI optimal criteria

        if pe and pe <= criteria['pe_max']:

            if roe_pct >= criteria['roe_min']:

                if debt_pct <= criteria['debt_equity_max']:

                    # Mark as AI OPTIMAL

                    r['signal'] = 'AI BUY'

                    filtered.append(r)

    

    # Sort by score and take top N

    filtered.sort(key=lambda x: x.get('score', 0), reverse=True)

    top_n = filtered[:criteria['max_positions']]

    

    # Re-rank within top N

    for i, r in enumerate(top_n):

        r['ai_rank'] = i + 1

    

    log_screener.info(f"AI Optimal: {len(top_n)} stocks match criteria (from {len(results)} scanned)")

    

    return top_n





def has_eod_cache():

    """Check if we have EOD cache data available"""

    fund_path = os.path.join(CACHE_DIR, 'fundamentals')

    if os.path.exists(fund_path):

        files = [f for f in os.listdir(fund_path) if f.endswith('.json')]

        return len(files) > 20  # At least 20 cached stocks

    return False





def safe_float(value, default=None):

    """Safely convert a value to float"""

    if value is None:

        return default

    try:

        return float(value)

    except (ValueError, TypeError):

        return default





def run_screener_eod(scope='france'):

    """Run screener using EOD Historical Data"""

    results = []

    

    # Get universe

    if scope == 'europe':

        exchanges = ['PA', 'AS', 'BR', 'MI', 'MC', 'XETRA']

    else:

        exchanges = ['PA']

    

    # Load tickers from universe cache or API

    all_tickers = []

    for exchange in exchanges:

        tickers = load_universe_from_eod(exchange)

        all_tickers.extend(tickers)

    

    # If no tickers loaded, try from fundamentals cache

    if not all_tickers:

        fund_path = os.path.join(CACHE_DIR, 'fundamentals')

        if os.path.exists(fund_path):

            for f in os.listdir(fund_path):

                if f.endswith('.json'):

                    ticker = f.replace('.json', '').replace('_', '.')

                    all_tickers.append({'ticker': ticker, 'name': '', 'exchange': ''})

    

    log_screener.info(f"Scanning {len(all_tickers)} tickers...")



    # Limit based on scope to avoid too long processing

    # France: 300 tickers (Euronext Paris)

    # Europe: 1000 tickers (6 exchanges)

    processed = 0

    max_tickers = 1000 if scope == 'europe' else 300



    log_screener.info(f"Processing up to {max_tickers} tickers for {scope}...")



    for stock in all_tickers:

        if processed >= max_tickers:

            log_screener.info(f"Reached limit of {max_tickers} tickers")

            break

            

        ticker = stock['ticker']

        

        # Get fundamentals from cache or API

        fund, err = eod_get_fundamentals(ticker, use_cache=True)

        if not fund or err:

            continue

        

        try:

            # Extract data

            general = fund.get('General', {})

            highlights = fund.get('Highlights', {})

            valuation = fund.get('Valuation', {})

            

            # Get latest financials

            financials = fund.get('Financials', {})

            income = financials.get('Income_Statement', {}).get('yearly', {})

            balance = financials.get('Balance_Sheet', {}).get('yearly', {})

            

            # Get most recent year data

            latest_year = None

            if income:

                years = sorted(income.keys(), reverse=True)

                if years:

                    latest_year = years[0]

            

            # Extract metrics with safe conversion

            pe = safe_float(highlights.get('PERatio')) or safe_float(valuation.get('TrailingPE'))

            roe = safe_float(highlights.get('ReturnOnEquityTTM'))

            profit_margin = safe_float(highlights.get('ProfitMargin'))

            

            # Get debt/equity from balance sheet

            debt_equity = None

            if latest_year and latest_year in balance:

                bs = balance[latest_year]

                total_debt = safe_float(bs.get('shortLongTermDebtTotal')) or safe_float(bs.get('longTermDebt')) or 0

                equity = safe_float(bs.get('totalStockholderEquity')) or safe_float(bs.get('totalEquity'))

                if equity and equity > 0:

                    debt_equity = total_debt / equity

            

            # Get current price and market cap

            market_cap = safe_float(highlights.get('MarketCapitalization'))

            price = None

            if market_cap:

                shares = safe_float(general.get('SharesOutstanding'))

                if shares and shares > 0:

                    price = market_cap / shares



            # Calculate 12-month momentum

            momentum_12m = None

            try:

                end_date = datetime.now().strftime('%Y-%m-%d')

                start_date = (datetime.now() - timedelta(days=380)).strftime('%Y-%m-%d')

                prices, err = eod_get_historical_prices(ticker, start_date, end_date, use_cache=True)



                if prices and len(prices) >= 2:

                    # Get oldest and newest price

                    oldest_price = safe_float(prices[0].get('close'))

                    newest_price = safe_float(prices[-1].get('close'))



                    if oldest_price and oldest_price > 0 and newest_price:

                        momentum_12m = (newest_price - oldest_price) / oldest_price

            except Exception:
                pass  # If momentum calculation fails, leave it as None



            # Calculate Higgons score

            score = calculate_higgons_score_for_screener(pe, roe, debt_equity, profit_margin)



            # Determine signal

            signal = determine_signal(pe, roe, score)



            # Determine country from exchange

            exchange = general.get('Exchange', stock.get('exchange', ''))

            country = get_country_from_exchange(exchange)



            results.append({

                'ticker': ticker,

                'name': general.get('Name', stock.get('name', ticker)),

                'country': country,

                'sector': general.get('Sector', 'Unknown'),

                'price': price,

                'market_cap': market_cap,

                'pe': pe,

                'roe': roe,

                'debt_equity': debt_equity,

                'profit_margin': profit_margin,

                'momentum_12m': momentum_12m,

                'score': score,

                'signal': signal,

                'source': 'EOD'

            })



            processed += 1



            # Debug log for first few tickers

            if processed <= 3:

                log_screener.debug(f"{ticker}: market_cap={market_cap}, momentum_12m={momentum_12m}")



            if processed % 50 == 0:

                log_screener.info(f"Processed {processed} tickers...")

                

        except Exception as e:

            log_screener.error(f"Error processing {ticker}: {e}")

            continue

    

    # Sort by score

    results.sort(key=lambda x: x.get('score', 0), reverse=True)

    

    log_screener.info(f"Found {len(results)} stocks with valid data")

    return results





def run_screener_yahoo():

    """Run screener using Yahoo Finance (legacy method)"""

    if not YFINANCE_OK:

        return []



    results = []

    for s in EUROPE_DB:

        r = s.copy()

        try:

            ticker_obj = yf.Ticker(s['ticker'])

            info = ticker_obj.info

            r['price'] = info.get('currentPrice') or info.get('regularMarketPrice')

            r['market_cap'] = info.get('marketCap')

            r['pe'] = info.get('trailingPE')

            r['roe'] = info.get('returnOnEquity')

            r['debt_equity'] = info.get('debtToEquity')

            if r['debt_equity']:

                r['debt_equity'] = r['debt_equity'] / 100  # Yahoo returns as percentage

            r['profit_margin'] = info.get('profitMargins')



            # Calculate 12-month momentum

            try:

                hist = ticker_obj.history(period='1y')

                if len(hist) >= 2:

                    oldest_price = hist['Close'].iloc[0]

                    newest_price = hist['Close'].iloc[-1]

                    if oldest_price > 0:

                        r['momentum_12m'] = (newest_price - oldest_price) / oldest_price

                    else:

                        r['momentum_12m'] = None

                else:

                    r['momentum_12m'] = None

            except Exception:
                r['momentum_12m'] = None



            r['score'] = calculate_higgons_score_for_screener(r['pe'], r['roe'], r['debt_equity'], r['profit_margin'])

            r['signal'] = determine_signal(r['pe'], r['roe'], r['score'])

            r['source'] = 'Yahoo'

        except Exception:
            r['score'] = 0

            r['signal'] = 'NEUTRE'

            r['source'] = 'Yahoo'

            r['market_cap'] = None

            r['momentum_12m'] = None

        results.append(r)



    results.sort(key=lambda x: x.get('score', 0), reverse=True)

    return results





def refresh_screener_data_background(scope='france'):

    """Refresh screener data in background by forcing API calls"""

    global REFRESH_STATUS



    REFRESH_STATUS['running'] = True

    REFRESH_STATUS['progress'] = 0

    REFRESH_STATUS['message'] = 'Initializing...'



    try:

        # Get universe

        if scope == 'europe':

            exchanges = ['PA', 'AS', 'BR', 'MI', 'MC', 'XETRA']

        else:

            exchanges = ['PA']



        # Load tickers from universe cache

        all_tickers = []

        for exchange in exchanges:

            tickers = load_universe_from_eod(exchange)

            all_tickers.extend(tickers)



        log_cache.info(f"REFRESH: Total universe: {len(all_tickers)} tickers for {scope}")



        # Filter tickers that need refresh (older than 7 days)

        tickers_to_refresh = []

        skipped = 0



        for stock in all_tickers:

            ticker = stock['ticker']



            # Check if fundamentals are fresh (< 7 days old)

            fund_cache_file = get_cache_path('fundamentals', ticker)

            needs_refresh = True



            if os.path.exists(fund_cache_file):

                try:

                    if is_cache_valid(fund_cache_file, 7):  # 7 days

                        needs_refresh = False

                        skipped += 1

                except Exception:
                    pass



            if needs_refresh:

                tickers_to_refresh.append(stock)



        log_cache.info(f"REFRESH: {len(tickers_to_refresh)} tickers need refresh, {skipped} already up-to-date")



        # Limit to reasonable number for Europe (500 max)

        max_refresh = 500 if scope == 'europe' else 300

        tickers_to_refresh = tickers_to_refresh[:max_refresh]



        REFRESH_STATUS['total'] = len(tickers_to_refresh)

        REFRESH_STATUS['message'] = f'Refreshing {len(tickers_to_refresh)} tickers ({skipped} skipped)...'



        log_cache.info(f"REFRESH: Starting PARALLEL refresh of {len(tickers_to_refresh)} tickers...")

        # Use ParallelAPIClient for faster fetching (10 workers, rate limited to 5/s for EOD API)
        client = ParallelAPIClient(max_workers=10, rate_limit=5.0)

        # Extract ticker strings for batch operations
        ticker_list = [stock['ticker'] for stock in tickers_to_refresh]

        # Batch fetch fundamentals in parallel
        REFRESH_STATUS['message'] = f'Fetching fundamentals ({len(ticker_list)} tickers)...'
        log_cache.info("REFRESH: Fetching fundamentals in parallel...")

        fund_results = client.fetch_fundamentals_batch(
            tickers=ticker_list,
            fetch_func=eod_get_fundamentals,
            use_cache=True,
            force_refresh=True
        )

        successful_funds = sum(1 for r in fund_results.values() if r.success)
        REFRESH_STATUS['progress'] = len(ticker_list) // 2  # 50% progress after fundamentals
        log_cache.info(f"REFRESH: Fundamentals complete: {successful_funds}/{len(ticker_list)} successful")

        # Batch fetch prices in parallel
        if REFRESH_STATUS['running']:
            REFRESH_STATUS['message'] = f'Fetching prices ({len(ticker_list)} tickers)...'
            log_cache.info("REFRESH: Fetching prices in parallel...")

            end_date = datetime.now().strftime('%Y-%m-%d')
            start_date = (datetime.now() - timedelta(days=380)).strftime('%Y-%m-%d')

            price_results = client.fetch_prices_batch(
                tickers=ticker_list,
                start_date=start_date,
                end_date=end_date,
                fetch_func=eod_get_historical_prices,
                use_cache=True
            )

            successful_prices = sum(1 for r in price_results.values() if r.success)
            REFRESH_STATUS['progress'] = len(ticker_list)  # 100% progress
            log_cache.info(f"REFRESH: Prices complete: {successful_prices}/{len(ticker_list)} successful")



        REFRESH_STATUS['message'] = 'Refresh complete! Running screener...'



        # Delete old screener caches to force full regeneration

        cache_base = CONFIG['screener_cache_file'].replace('.json', '')

        for cache_suffix in ['_default.json', f'_{scope}_standard.json', f'_{scope}_ai_optimal.json']:

            cache_file = cache_base + cache_suffix

            if os.path.exists(cache_file):

                os.remove(cache_file)

                log_cache.info(f"REFRESH: Deleted cache: {os.path.basename(cache_file)}")



        # Re-run screener to update cache

        log_cache.info("REFRESH: Regenerating screener cache with new data...")

        run_screener(force=True, scope=scope, mode='standard')



        REFRESH_STATUS['message'] = 'Done!'

        log_cache.info("REFRESH: Background refresh completed!")



    except Exception as e:

        REFRESH_STATUS['message'] = f'Error: {str(e)}'

        log_cache.error(f"REFRESH: Error during background refresh: {e}")

    finally:

        REFRESH_STATUS['running'] = False





def calculate_higgons_score_for_screener(pe: Optional[float], roe: Optional[float], debt_equity: Optional[float], profit_margin: Optional[float]) -> int:

    """Calculate Higgons-style score for screener"""

    score = 0

    

    # PE Score (0-35 points) - lower is better

    if pe and pe > 0:

        if pe < 7: score += 35

        elif pe < 10: score += 30

        elif pe < 12: score += 25

        elif pe < 15: score += 15

        elif pe < 20: score += 5

    

    # ROE Score (0-30 points) - higher is better

    if roe and roe > 0:

        roe_pct = roe if roe < 1 else roe / 100  # Handle both 0.15 and 15 formats

        if roe_pct >= 0.25: score += 30

        elif roe_pct >= 0.20: score += 25

        elif roe_pct >= 0.15: score += 20

        elif roe_pct >= 0.12: score += 15

        elif roe_pct >= 0.10: score += 10

    

    # Debt/Equity Score (0-20 points) - lower is better

    if debt_equity is not None:

        if debt_equity <= 0: score += 20  # Net cash

        elif debt_equity <= 0.3: score += 15

        elif debt_equity <= 0.5: score += 10

        elif debt_equity <= 1.0: score += 5

    

    # Profit Margin Score (0-15 points) - higher is better

    if profit_margin and profit_margin > 0:

        pm = profit_margin if profit_margin < 1 else profit_margin / 100

        if pm >= 0.15: score += 15

        elif pm >= 0.10: score += 10

        elif pm >= 0.05: score += 5

    

    return score





def determine_signal(pe, roe, score):

    """Determine investment signal based on criteria"""

    if pe and roe:

        roe_val = roe if roe < 1 else roe / 100

        

        # ACHAT: PE <= 12 AND ROE >= 10% AND score >= 50

        if pe <= 12 and roe_val >= 0.10 and score >= 50:

            return 'ACHAT'

        

        # WATCH: PE <= 15 AND ROE >= 8%

        if pe <= 15 and roe_val >= 0.08:

            return 'WATCH'

        

        # CHER: PE > 20 OR ROE < 5%

        if pe > 20 or roe_val < 0.05:

            return 'CHER'

    

    return 'NEUTRE'





def get_country_from_exchange(exchange):

    """Get country code from exchange"""

    exchange_map = {

        'PA': 'FR', 'XPAR': 'FR', 'Paris': 'FR',

        'AS': 'NL', 'XAMS': 'NL', 'Amsterdam': 'NL',

        'BR': 'BE', 'XBRU': 'BE', 'Brussels': 'BE',

        'MI': 'IT', 'XMIL': 'IT', 'Milan': 'IT',

        'MC': 'ES', 'XMAD': 'ES', 'Madrid': 'ES',

        'XETRA': 'DE', 'XFRA': 'DE', 'Frankfurt': 'DE',

        'LSE': 'UK', 'XLON': 'UK', 'London': 'UK',

        'SW': 'CH', 'XSWX': 'CH', 'Swiss': 'CH',

        'VI': 'AT', 'XWBO': 'AT', 'Vienna': 'AT',

        'HE': 'FI', 'XHEL': 'FI', 'Helsinki': 'FI',

        'ST': 'SE', 'XSTO': 'SE', 'Stockholm': 'SE',

        'OL': 'NO', 'XOSL': 'NO', 'Oslo': 'NO',

        'CO': 'DK', 'XCSE': 'DK', 'Copenhagen': 'DK',

        'LS': 'PT', 'XLIS': 'PT', 'Lisbon': 'PT',

        'AT': 'GR', 'XATH': 'GR', 'Athens': 'GR',

    }

    

    for key, country in exchange_map.items():

        if key.lower() in exchange.lower():

            return country

    

    return 'EU'



def get_security_data(ticker: str) -> Dict[str, Any]:

    """Get detailed security data for the detail page"""

    data = {'ticker': ticker, 'name': '', 'sector': '', 'country': '', 'price': 0, 'change': 0, 'change_pct': 0,

            'pe': None, 'forward_pe': None, 'roe': None, 'market_cap': None, 'dividend_yield': None, 'beta': None,

            'high_52w': None, 'low_52w': None, 'volume': None, 'avg_volume': None, 'eps': None, 'revenue': None,

            'profit_margin': None, 'debt_equity': None, 'current_ratio': None, 'book_value': None, 'target_price': None,

            'description': '', 'industry': '', 'website': '', 'employees': None, 'memo_file': None, 'memo_content': None,

            'price_history': [], 'volume_history': []}

    

    # Find in DB (try with and without .PA)

    for s in EUROPE_DB:

        if s['ticker'] == ticker or s['ticker'] == ticker + '.PA' or s['ticker'].replace('.PA', '') == ticker:

            data.update({k: s[k] for k in ['name', 'sector', 'country']})

            break

    

    # Find memo file and read content

    data['memo_file'] = find_memo(ticker)

    if data['memo_file']:

        data['memo_content'] = read_memo_content(data['memo_file'])

    

    if YFINANCE_OK:

        try:

            yf_ticker = get_yf_ticker(ticker)

            info = yf.Ticker(yf_ticker).info

            data.update({

                'name': info.get('shortName') or info.get('longName') or data['name'],

                'price': info.get('currentPrice') or info.get('regularMarketPrice') or 0,

                'change': info.get('regularMarketChange') or 0,

                'change_pct': info.get('regularMarketChangePercent') or 0,

                'pe': info.get('trailingPE'), 'forward_pe': info.get('forwardPE'),

                'roe': info.get('returnOnEquity'), 'market_cap': info.get('marketCap'),

                'dividend_yield': info.get('dividendYield'), 'beta': info.get('beta'),

                'high_52w': info.get('fiftyTwoWeekHigh'), 'low_52w': info.get('fiftyTwoWeekLow'),

                'volume': info.get('volume'), 'avg_volume': info.get('averageVolume'),

                'eps': info.get('trailingEps'), 'revenue': info.get('totalRevenue'),

                'profit_margin': info.get('profitMargins'), 'debt_equity': info.get('debtToEquity'),

                'current_ratio': info.get('currentRatio'), 'book_value': info.get('bookValue'),

                'target_price': info.get('targetMeanPrice'),

                'description': info.get('longBusinessSummary') or '',

                'sector': info.get('sector') or data['sector'],

                'industry': info.get('industry') or '',

                'website': info.get('website') or '',

                'employees': info.get('fullTimeEmployees')

            })

            

            # Get price history (1 year)

            try:

                t = yf.Ticker(yf_ticker)

                hist = t.history(period="1y")

                if not hist.empty:

                    data['price_history'] = [

                        {'date': d.strftime('%Y-%m-%d'), 'close': round(row['Close'], 2), 'volume': int(row['Volume'])}

                        for d, row in hist.iterrows()

                    ]

            except Exception as e:

                log_api.error(f"Error fetching history for {ticker}: {e}")

                

        except Exception as e:

            log_api.error(f"Error fetching {ticker}: {e}")

    return data
